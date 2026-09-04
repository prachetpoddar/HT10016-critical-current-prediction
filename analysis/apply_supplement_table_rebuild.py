"""
apply_supplement_table_rebuild.py

Replace the printed Tables S4, S5 and S6 of the Supplemental Material with the
rows analysis/build_supplement_tables.py generates from the deposit, and rewrite
the prose that describes them.

Why. The printed tables were maintained by hand and had drifted from the files
they say they excerpt.

  * Table S4 printed three rows from papers in audit/withdrawn_records.csv, one
    from physc.2010.03.003 and two from S0011-2275(97)00151-3. The anchor file
    holds none of them.
  * Table S5 printed three, the same two papers plus 0921-4534(94)00021-2. The
    field-axis fit file holds no rows for any of them.
  * Table S6 printed the three Co0.05Fe0.95Se rows at 4.2 K as delivered
    predictions with the refusal column reading "none". In the deposit those
    rows carry predicted_log_Jc of NaN with the refusal flags
    H_below_validated_reduced_field, H_below_validated_reduced_field and
    family_fails_field_axis_validation; the printed numbers are the
    withheld_log_Jc column, which is what the pipeline computed and then
    declined to deliver.

The prose is rewritten with the tables because it describes rows by name. The
generator's selection rule now also keeps the fit whose resolved scale sits
furthest below its literature value, so the failure mode Sec. III.F quantifies
stays in the table rather than being dropped by a rule that sorts on span.

One claim in the Table S6 paragraph is corrected in the other direction. It said
a row carries either a prediction or a refusal code and no value. A refusal acts
on one prediction target rather than on a candidate, so a row refused on the
field axis can still carry the temperature-axis value the remaining gates allow,
and 321 of the 540 rows refused for a missing upper-critical-field anchor do.

Usage:
    python3 analysis/apply_supplement_table_rebuild.py --supp S.docx \
        --out-dir DIR [--dry-run]
"""
import argparse
import copy
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import docx  # noqa: E402
import pandas as pd  # noqa: E402

import build_supplement_tables as gen  # noqa: E402
from apply_manuscript_edits import apply  # noqa: E402

# table index in the .docx -> generator key, identified by its header row so a
# reordered document cannot silently write S6's rows into S5
TARGETS = {
    "S4": ("Source paper", "Compound", "Form", "Sample", "Tc",
           "Anchor point", "log10 Jc", "n"),
    "S5": ("Source paper", "Compound", "Hc2 used / default (T)",
           "Scale provenance", "Span", "βH (SE)", "Flag"),
    "S6": ("Candidate", "Anchors Tc / Hc2", "Scope", "T (K)", "H (T)",
           "log10 Jc", "95% interval", "Refusal"),
}


def find_table(doc, header):
    """Locate a table by its header row, not by position."""
    want = [h.strip().lower() for h in header]
    for t in doc.tables:
        got = [c.text.strip().lower() for c in t.rows[0].cells]
        if got == want:
            return t
    return None


def set_cell(cell, text):
    """Write text keeping the cell's first run, so the table's formatting and
    fonts survive. Everything after the first run and the first paragraph is
    removed, because a stale second run leaves the old value visible."""
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    par = cell.paragraphs[0]
    if not par.runs:
        par.add_run("")
    par.runs[0].text = text
    for r in par.runs[1:]:
        r._element.getparent().remove(r._element)


def resize(table, n_data):
    """Make the table hold exactly n_data rows below the header, cloning the
    last data row so a new row inherits its borders and shading."""
    while len(table.rows) - 1 > n_data:
        row = table.rows[-1]._element
        row.getparent().remove(row)
    while len(table.rows) - 1 < n_data:
        new = copy.deepcopy(table.rows[-1]._element)
        table.rows[-1]._element.addnext(new)


def rewrite(doc, key, rows, report):
    t = find_table(doc, TARGETS[key])
    if t is None:
        report.append((key, False, "header row not found", "", ""))
        return False
    if len(t.columns) != len(rows[0]):
        report.append((key, False, "%d columns in the document against %d "
                       "generated" % (len(t.columns), len(rows[0])), "", ""))
        return False
    before = len(t.rows) - 1
    resize(t, len(rows))
    for r, values in zip(t.rows[1:], rows):
        for cell, v in zip(r.cells, values.values()):
            set_cell(cell, str(v))
    report.append((key, True, "%d data row(s) -> %d" % (before, len(rows)),
                   "", ""))
    return True


def prose(rows):
    """The sentences that name rows by name, rebuilt from the rows written."""
    s5 = rows["S5"]
    ok = [r for r in s5 if r["flag"] == "ok"]
    bound = [r for r in s5 if r["flag"] == "bound"]
    ceiling = [r for r in bound if r["beta"].startswith("30.00")]
    worst = max(bound, key=lambda r: float(r["hc2"].split("/")[1])
                / float(r["hc2"].split("/")[0]))
    se_over = sum(1 for r in ceiling
                  if float(r["beta"].split("(")[1].rstrip(")")) > 30.0)
    word = {1: "one", 2: "two", 3: "three", 4: "four", 5: "five",
            6: "six", 7: "seven"}
    p = pd.read_csv(gen.PRED, low_memory=False)
    hc2_un = p[p.refusal_flag == "Hc2_unavailable"]
    carried = int(hc2_un.predicted_log_Jc.notna().sum())

    s5_prose = (
        "These rows are selected by rule rather than by hand, and the rule is "
        "in the generator named in the caption. The %s passing fits are the "
        "smallest standard errors in the cohort, one per source paper, with "
        "spans of %s. The %s cuprate rows are the narrowest spans in the "
        "cohort, %s to %s, and all of them sit at the imposed ceiling of 30, "
        "%s with a standard error larger than the estimate. The %s row is the "
        "failure mode quantified in Sec. III.F: the resolved scale of %s T "
        "sits far below the %s T literature value, the span is %s, and the "
        "exponent runs to %s. These are the fits the applicability "
        "qualification exists to exclude, and they are visible in the deposit "
        "rather than filtered out of it."
        % (word.get(len(ok), str(len(ok))),
           ", ".join(r["span"] for r in ok),
           word.get(len(ceiling), str(len(ceiling))),
           ceiling[0]["span"], ceiling[-1]["span"],
           word.get(se_over, str(se_over)),
           worst["compound"],
           ("%g" % float(worst["hc2"].split("/")[0])),
           ("%g" % float(worst["hc2"].split("/")[1])),
           worst["span"], worst["beta"].split(" (")[0]))

    s6_prose = (
        "Each row is one candidate compound at one point of the evaluation "
        "grid. A row carries a prediction with a bootstrap interval, a refusal "
        "code, or both: a refusal acts on one prediction target rather than on "
        "a candidate, so a row refused on the field axis can still carry the "
        "temperature-axis value the remaining gates allow, and %d of the %d "
        "rows refused for a missing upper-critical-field anchor do."
        % (carried, len(hc2_un)))
    return s5_prose, s6_prose


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--supp", required=True)
    ap.add_argument("--out-dir", default=".")
    ap.add_argument("--dry-run", action="store_true")
    args = ap.parse_args()
    if not os.path.isdir(gen.DATA):
        sys.exit("run from the repository root")

    banned = gen.withdrawn_tokens()
    rows = {"S4": gen.table_s4(banned), "S5": gen.table_s5(banned),
            "S6": gen.table_s6(banned)}
    doc = docx.Document(args.supp)

    report = []
    ok = all(rewrite(doc, k, rows[k], report) for k in ("S4", "S5", "S6"))

    s5_prose, s6_prose = prose(rows)
    edits = [
        ("These rows were selected to be useful rather than favourable. The "
         "Nb3Sn wire rows are the well-behaved case: a directly matched "
         "critical field at the measurement temperature, a span of 0.79, and "
         "an exponent near 1.1 with a small standard error. The Ba(Fe,Ru)2As2 "
         "row is the failure mode quantified in Sec. III.F: the resolved scale "
         "of 15.3 T sits far below the 60 T literature value, the resulting "
         "span is 0.033, and the exponent runs to 19.1, well outside the range "
         "this fit can constrain. The two cuprate rows sit at the bound of 30 "
         "with standard errors larger than the estimates. These are the fits "
         "the anchor-count gate and the applicability qualification exist to "
         "exclude, and they are visible in the deposit rather than filtered "
         "out of it.",
         s5_prose, "named rows that the fit file no longer holds", None),

        ("Each row is one candidate compound at one point of the evaluation "
         "grid. A row either carries a prediction with a bootstrap interval or "
         "carries a refusal code and no value.",
         s6_prose,
         "321 of the 540 Hc2-unavailable rows carry both", None),

        ("Table S6. Excerpt from the candidate dispatch file, showing one "
         "dispatched compound and one refused compound across the grid. Both "
         "are iron chalcogenide 11-type and iron pnictide 122-type "
         "respectively.",
         "Table S6. Excerpt from the candidate dispatch file, generated by "
         "analysis/build_supplement_tables.py: two dispatched targets, then "
         "one target carrying each refusal code the file uses. The dispatched "
         "rows are MgB2-class, which is the only family the routine "
         "dispatches; the refused rows are iron chalcogenide 11-type and iron "
         "pnictide 122-type.",
         "the caption named the wrong families and the wrong selection rule",
         None),

        ("Table S5. Excerpt from the field-axis Form 3 fit file.",
         "Table S5. Excerpt from the field-axis Form 3 fit file, generated by "
         "analysis/build_supplement_tables.py.",
         "name the generator, since the prose now points at it", None),

        ("The anchor is the measured critical current density at the lowest "
         "temperature and field at which that sample was measured, so it "
         "involves no fitted exponent and no critical-field scale.",
         "Each row records one isotherm's anchor point, the temperature and "
         "field at which that measurement was taken, so the anchor involves "
         "no fitted exponent and no critical-field scale.",
         "false for a multi-isotherm sample: physc.2009.05.098 contributes "
         "nine rows from 2 K to 40 K, and only one of them is at the lowest "
         "temperature that sample was measured at", None),

        ("Table S4. Excerpt from the per-paper critical-current anchor file.",
         "Table S4. Excerpt from the per-paper critical-current anchor file, "
         "generated by analysis/build_supplement_tables.py.",
         "name the generator", None),
    ]
    misses = apply(doc, edits, "supplement", report)

    for label, done, a, b, why in report:
        print("   %-11s %-4s %s" % (label, "ok" if done else "MISS", a[:74]))
    print()
    if misses or not ok:
        sys.exit("%d edit(s) did not match; nothing written"
                 % (len(misses) + (0 if ok else 1)))
    if args.dry_run:
        print("all matched; --dry-run, nothing written")
        return 0
    os.makedirs(args.out_dir, exist_ok=True)
    out = os.path.join(args.out_dir, os.path.basename(args.supp))
    doc.save(out)
    print("written %s" % out)
    return 0


if __name__ == "__main__":
    sys.exit(main())
