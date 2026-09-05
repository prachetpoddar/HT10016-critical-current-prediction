#!/usr/bin/env python3
"""
apply_repaired_cohort_edits.py

Put the repaired cohort into the manuscript, the supplement and the response.

Every edit is a literal find-and-replace declared below with the deposited file
its new value comes from, and the script refuses to write anything if any edit
fails to find its target. That rule is inherited from
analysis/apply_manuscript_edits.py and it is why this deposit no longer silently
disagrees with its own manuscript.

WHAT IS AND IS NOT EDITED HERE. Only counts whose value is a recount are
applied. A number quoted alongside a statistic computed on that cohort is NOT
renumbered, because changing 94 to 52 beside a median computed over the 94 would
be worse than leaving it. Those passages are printed as blockers, and the
document is not finished until they are cleared.

    python3 analysis/apply_repaired_cohort_edits.py --dry-run
    python3 analysis/apply_repaired_cohort_edits.py --out-dir <dir>

Run from the repository root.
"""
import argparse
import copy
import os
import shutil
import sys

import docx

SRC = "/mnt/user-data/uploads/SuperconductorWorkflow"
MS = "HT10016_revised_final.docx"
SUPP = "SUPPLEMENTAL_MATERIAL_revised_final.docx"
RESP = "RESPONSE_TO_REFEREES_final.docx"

# (find, replace, why)
MS_EDITS = [
    ("built from 62 papers that contribute fitted critical-current curves, "
     "covering 38 compounds and 4146 extracted data points",
     "built from 50 papers that contribute fitted critical-current curves, "
     "covering 35 compounds and 3303 extracted data points",
     "abstract; provenance_table_fitcohort_full.csv, status == contributing"),
    ("Sixty-two papers pass the fittability filters and contribute fitted "
     "curves across 38 compounds and 4146 critical-current data points",
     "Fifty papers pass the fittability filters and contribute fitted "
     "curves across 35 compounds and 3303 critical-current data points",
     "Fig. 1 caption; the same source"),
]
MS_TABLE_EDITS = [
    (0, 2, 1, "62", "50", "Table I, papers contributing fitted curves"),
    (0, 3, 1, "38", "35", "Table I, distinct compounds"),
    (0, 4, 1, "4146", "3303", "Table I, points extracted"),
    (0, 6, 1, "260", "257", "Table I, temperature-axis fits; "
                            "phase_3_p44_post_UCLA_beta_T_fits_repaired.csv"),
    (0, 7, 1, "94", "52", "Table I, field-axis fits passing; "
                          "audit/fit_protocol_applied.csv"),
    (0, 7, 2, "Field-exponent aggregation, from 16 source papers",
     "Field-exponent aggregation, from 12 source papers",
     "Table I, field-axis source papers"),
]
SUPP_EDITS = [
    ("covers the 62 source papers that contribute fitted data",
     "covers the 50 source papers that contribute fitted data",
     "Sec. 11; provenance_table_fitcohort_full.csv, status == contributing"),
    ("The 62 papers listed contribute curves",
     "The 50 papers listed contribute curves",
     "Table S1 caption; the same source"),
]
RESP_EDITS = [
    ("934 articles screened, 62 contributing fitted curves, 38 distinct "
     "compounds, 4146 extracted critical-current points, 23 fully fittable "
     "compounds, 260 temperature-axis fits, 94 field-axis fits drawn from 16 "
     "source papers, and 96 per-paper anchors behind Figure 3.",
     "934 articles screened, 50 contributing fitted curves, 35 distinct "
     "compounds, 3303 extracted critical-current points, 257 temperature-axis "
     "fits, 52 field-axis fits drawn from 12 source papers, and 96 per-paper "
     "anchors behind Figure 3. The counts for contributing papers, compounds "
     "and extracted points are lower than in the previous revision because "
     "eleven papers were withdrawn from the cohort and their provenance rows "
     "had not been removed with them; the reasons are set out below.",
     "the Table I ladder, restated on the repaired cohort"),
    ("The figure of 62 is no longer introduced first in the Supplemental "
     "Material", "The figure of 50 is no longer introduced first in the "
     "Supplemental Material", "the same sentence, following the ladder"),
]

# Passages that quote one of these cohorts ALONGSIDE a statistic computed on it.
# Renumbering them without recomputing the statistic would put the document back
# into the state this whole exercise exists to get it out of.
BLOCKERS = [
    (MS, "Using the 260 temperature-axis fits",
     "quotes 260 with per-family counts and bootstrap fractions computed on "
     "it; the three fits lost are two chalcogenide and one 122"),
    (MS, "the 260-fit temperature-exponent cohort contains no MgB2 fits",
     "the statement stays true at 257, but the number has to move with the "
     "recomputation above"),
    (MS, "for 15 of 94 curves it exceeds 0.9",
     "the scale audit's own ratio statistic, computed over the field-axis "
     "cohort; needs re-running on 52, not renumbering"),
    (SUPP, "for 15 of 94 curves it exceeds 0.9", "the same statistic"),
    (SUPP, "gives 1.158 over 94 fits from the same 16 papers",
     "a pooled median and bootstrap interval over the 94; both move"),
    (RESP, "15 of 94 curves sit above 0.9", "the same statistic"),
    (SUPP, "among the 23 compounds whose per-compound aggregate Form 3 fit "
           "converges",
     "23 is defined here as the compounds whose per-compound aggregate fit "
     "converges, computed by run_closed_form_fits.py on a dataset that still "
     "contains the eleven withdrawn papers; it has to be re-derived, not "
     "renumbered, and the Table I row that carries it is left alone until it is"),
]

INSERT_AFTER = "We attempted a rebuild against a properly temperature-resolved"
INSERT = os.path.join("audit", "disclosure_draft_20260905.md")


def replace_in_paragraph(p, find, repl):
    """Replace across runs, keeping the first run's formatting."""
    text = "".join(r.text for r in p.runs)
    if find not in text:
        return False
    start = text.index(find)
    end = start + len(find)
    pos, first, spans = 0, None, []
    for r in p.runs:
        a, b = pos, pos + len(r.text)
        if b > start and a < end:
            spans.append((r, max(start, a) - a, min(end, b) - a))
            if first is None:
                first = r
        pos = b
    if first is None:
        return False
    for r, i, j in spans:
        r.text = r.text[:i] + (repl if r is first else "") + r.text[j:]
    return True


def apply_doc(path, edits, table_edits=()):
    d = docx.Document(path)
    missed = []
    for find, repl, why in edits:
        if not any(replace_in_paragraph(p, find, repl) for p in d.paragraphs):
            missed.append((why, find[:60]))
    for ti, row, col, find, repl, why in table_edits:
        cell = d.tables[ti].rows[row].cells[col]
        if not any(replace_in_paragraph(p, find, repl) for p in cell.paragraphs):
            missed.append((why, find[:60]))
    return d, missed


def insert_disclosure(d):
    """Insert the disclosure paragraphs after the located anchor paragraph."""
    anchor = None
    for p in d.paragraphs:
        if p.text.startswith(INSERT_AFTER):
            anchor = p
    if anchor is None:
        return 0
    body = open(INSERT).read()
    body = body.split("---", 2)[1] if body.count("---") >= 2 else body
    body = body.split("# Edits to passages already in the response")[0]
    blocks = [b.strip() for b in body.split("\n\n") if b.strip()]
    at = anchor
    n = 0
    for b in blocks:
        txt = " ".join(b.split())
        if txt.startswith("#"):
            txt = txt.lstrip("# ").strip()
        txt = txt.replace("**", "")
        new = copy.deepcopy(anchor._p)
        at._p.addnext(new)
        para = docx.text.paragraph.Paragraph(new, anchor._parent)
        for r in para.runs[1:]:
            r.text = ""
        if para.runs:
            para.runs[0].text = txt
        at = para
        n += 1
    return n


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--out-dir")
    ap.add_argument("--dry-run", action="store_true")
    a = ap.parse_args()

    docs = [(MS, MS_EDITS, MS_TABLE_EDITS), (SUPP, SUPP_EDITS, ()),
            (RESP, RESP_EDITS, ())]
    results, all_missed = [], []
    for name, edits, tedits in docs:
        d, missed = apply_doc(os.path.join(SRC, name), edits, tedits)
        n_ins = insert_disclosure(d) if name == RESP else 0
        results.append((name, d, len(edits) + len(tedits), n_ins))
        all_missed += [(name,) + m for m in missed]
        print(f"{name}: {len(edits) + len(tedits)} edits, "
              f"{len(missed)} not found" + (f", {n_ins} paragraphs inserted"
                                            if n_ins else ""))
    if all_missed:
        print()
        print("EDITS THAT FOUND NO TARGET. Nothing written.")
        for m in all_missed:
            print("  ", m)
        return 1

    print()
    print("BLOCKERS: passages that quote a changed cohort beside a statistic")
    print("computed on it. These are NOT renumbered here and the documents are")
    print("not finished until each is recomputed.")
    for doc, text, why in BLOCKERS:
        d = docx.Document(os.path.join(SRC, doc))
        found = any(text in p.text for p in d.paragraphs)
        print(f"  [{'found' if found else 'NOT FOUND'}] {doc}: {text[:52]}")
        print(f"           {why}")

    if a.dry_run or not a.out_dir:
        print()
        print("dry run: nothing written")
        return 0
    os.makedirs(a.out_dir, exist_ok=True)
    for name, d, _, _ in results:
        out = os.path.join(a.out_dir, name.replace("_final", "_repaired"))
        d.save(out)
        print(f"written: {out}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
