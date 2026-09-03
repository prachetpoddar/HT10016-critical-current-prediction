#!/usr/bin/env python3
"""Apply the audited corrections to the manuscript and the supplement.

Every edit is a literal find-and-replace declared in EDITS below, with the
deposited file each new value comes from. The script refuses to write anything
if any edit fails to find its target, because a partially applied correction is
how this deposit came to disagree with its own manuscript in the first place.

Replacement happens at run level. A Word paragraph stores formatting per run, so
rewriting a whole paragraph would flatten the subscripts in betaT, Hc2 and
log10 Jc. The matched span is located across runs, the new text is written into
the first run of the span in that run's formatting, and the rest of the span is
cleared. Text on either side keeps its own runs untouched.

    python analysis/apply_manuscript_edits.py --ms <docx> --supp <docx> --dry-run
    python analysis/apply_manuscript_edits.py --ms <docx> --supp <docx> --out-dir .
"""
import argparse
import os
import shutil
import sys

import docx


# (document, find, replace, why)
MS_EDITS = [
    # --- Table I ------------------------------------------------------------
    ("69", "62", "Table I, papers contributing fitted curves; "
                 "provenance_table_fitcohort_full.csv", "table:0:2:1"),
    ("43", "38", "Table I, distinct compounds; same source", "table:0:3:1"),
    ("4387", "4146", "Table I, points extracted; n_Jc_points summed", "table:0:4:1"),
    ("419", "260", "Table I, temperature-axis partial fits; "
                   "phase_3_p44_post_UCLA_beta_T_fits.csv", "table:0:6:1"),
    ("95", "94", "Table I, field-axis fits passing physicality; "
                 "phase_3_form3_fits_partial_cohortB_v2.csv", "table:0:7:1"),
    ("Field-exponent aggregation, from 17 source papers",
     "Field-exponent aggregation, from 16 source papers",
     "Table I, field-axis source papers; distinct arxiv_id among passing fits",
     "table:0:7:2"),
    ("110", "96", "Table I, per-paper anchors; "
                   "phase_3_p31_jc_anchor_per_paper.csv", "table:0:8:1"),
    ("185", "183", "Table I, candidate compounds; "
                   "phase_3_p56_candidate_tier_assignment.csv", "table:0:9:1"),

    # --- Sec. II.A, source composition ---------------------------------------
    ("Twenty-nine of the 33 source papers behind the temperature-axis exponent "
     "cohort are arXiv preprints",
     "Eighteen of the 20 source papers behind the temperature-axis exponent "
     "cohort are arXiv preprints",
     "20 distinct paper_id values in the temperature-axis fit table", None),

    # --- Fig. 3 caption -------------------------------------------------------
    ("Of the 110 per-paper anchor records of Table I, 61 fall in the three "
     "families shown",
     "Of the 96 per-paper anchor records of Table I, 52 fall in the three "
     "families shown",
     "anchor table filtered to the three plotted families", None),
    ("The 61 collapse to the 44 markers drawn",
     "The 52 collapse to the 34 markers drawn",
     "aggregate_per_physical_sample over those records", None),

    # --- Sec. III.C, temperature axis ----------------------------------------
    ("Using the 419 temperature-axis fits", "Using the 260 temperature-axis fits",
     "phase_3_p44_post_UCLA_beta_T_fits.csv", None),
    ("reaches 0.261 across five compounds and 37 fits, with 92% of bootstrap "
     "resamples below the screening-grade threshold of 1",
     "reaches 0.588 across five compounds and 89 fits, with every scorable "
     "bootstrap resample below the screening-grade threshold of 1",
     "compound_leave_one_out.py; audit/temperature_axis_leave_one_out.csv", None),
    ("Iron pnictide 122-type reaches 1.092 across three compounds, with only "
     "38% of resamples below threshold.",
     "Iron pnictide 122-type reaches 1.314 across two compounds, with 30% of "
     "resamples below threshold.", "same", None),
    ("Iron pnictide 1111-type reaches 1.721 across four compounds, with 8% "
     "below threshold.",
     "Iron pnictide 1111-type reaches 3.120 across three compounds, with 21% "
     "below threshold.", "same", None),

    # --- Sec. III.C, the 122 paragraph ---------------------------------------
    ("The 122-type result is dominated by a single compound. Holding out "
     "BaFe2As2, which supplies 147 of the 198 fits, gives 1.230; the other two "
     "compounds give 0.686 and 0.701. The family figure is therefore largely a "
     "statement about how well two compounds predict a third, which is the "
     "honest reading of a three-compound leave-one-out at this cohort size.",
     "The 122-type family now holds two compounds on this axis, Ba(FeAs)2 with "
     "85 fits and K(FeAs)2 with 21. A leave-one-compound-out across two "
     "compounds is a statement about how well each predicts the other, and we "
     "report it as that rather than as a family-level generalization test.",
     "phase_3_p44_post_UCLA_beta_T_fits.csv, compound counts per family", None),
    ("the 419-fit temperature-exponent cohort contains no MgB2 fits",
     "the 260-fit temperature-exponent cohort contains no MgB2 fits",
     "same", None),
    ("with a median βT of 1.14.",
     "with a median βT of 1.14; two of the fifteen are poorly constrained, one "
     "fitting with an rms of 6.3 in log10 Jc and one resting on three points.",
     "h1b_per_paper_form3_fits.csv, phase_f_filename column", None),

    # --- Sec. III.C, field axis ----------------------------------------------
    ("gives 0.641 for iron chalcogenide 11-type, 0.753 for MgB2-class, and "
     "0.973 for iron pnictide 122-type, with iron pnictide 1111-type at 3.07 "
     "at this cohort scope and 5.13 under the expanded compound-diversity "
     "validation.",
     "gives 0.753 for MgB2-class, 0.973 for iron pnictide 122-type and 1.093 "
     "for iron chalcogenide 11-type, with iron pnictide 1111-type at 2.571.",
     "data/phase_3_p47_compound_leave_out_MAE.csv", None),
    ("The ordering across families is the same on both axes, and it is not an "
     "artifact of the conditioning protocol: repeating the field-axis test with "
     "a substructure-median predictor, matching the temperature-axis protocol "
     "exactly, gives 0.558, 0.751, 0.929, and 3.13, preserving the ordering.",
     "The ordering across families is no longer the same on both axes: iron "
     "chalcogenide 11-type is first on the temperature axis and third on the "
     "field axis. Repeating the field-axis test with a substructure-median "
     "predictor, matching the temperature-axis protocol exactly, gives 0.751, "
     "0.929, 1.094 and 2.622, and reproduces the field-axis ordering.",
     "same", None),

    # --- Sec. III.C, the graded applicability claim ---------------------------
    ("Iron chalcogenide 11-type is the only family that passes on the "
     "temperature axis, whose critical scale is directly reported, and it "
     "passes on the field axis as well. Iron pnictide 122-type and MgB2-class "
     "pass on the field axis only, and their validation inherits the "
     "qualification of Section III.F.",
     "Iron chalcogenide 11-type is the only family that passes on the "
     "temperature axis, whose critical scale is directly reported, and it does "
     "not pass on the field axis. MgB2-class and iron pnictide 122-type pass on "
     "the field axis only, and their validation inherits the qualification of "
     "Section III.F. No family passes on both axes.",
     "data/phase_3_p47_compound_leave_out_MAE.csv", None),

    # --- Sec. III.B and the summary, the conditioning claim -------------------
    ("On the matched five-family cohort the reduction is 16-fold on means and "
     "4.7-fold on medians. All three figures are computed on \u03b2H and therefore "
     "inherit the field-scale qualification of Section III.F. We no longer "
     "present a single \"23-fold\" figure, and the conditioning claim below does "
     "not rest on it.",
     "Restricting to the matched five-family cohort removes that particular "
     "defect but not the more serious one: in both comparisons the predictor "
     "for a family is built from a pool that contains that family's own fits, "
     "so neither figure is a statement about generalization. We therefore "
     "report the comparison under leave-one-substructure-out, in which the "
     "held-out family is withheld at every stage. The improvement is then "
     "between one and about two-fold and depends on the cohort: across the "
     "seven families that carry a descriptor the best Stage 2 reading gives "
     "1.07-fold, restricted to fits passing physicality it gives 2.24-fold, "
     "and with the cuprate families removed 1.83-fold. Stage 3 gives 1.37-fold "
     "on means across all seven families and 2.20-fold with the cuprate "
     "families removed, where its interquartile bound covers the residual in "
     "four of four families. All figures are computed on \u03b2H and therefore "
     "inherit the field-scale qualification of Section III.F. We no longer "
     "present a single fold-improvement headline.",
     "analysis/multi_stage_loso.py; audit/multi_stage_loso.csv", None),
    ("Stage 1 gives a leave-one-substructure-out mean absolute error of 10.10 "
     "in the dimensionless field exponent.",
     "Stage 1 gives a leave-one-substructure-out mean absolute error of 12.30 "
     "in the dimensionless field exponent, across the seven substructure "
     "families that carry a descriptor.",
     "analysis/multi_stage_loso.py, cohort 'all fits, every family with a "
     "descriptor'", None),
    ("The Stage 1 error is computed across nine substructure families and the "
     "Stage 2 error across the five families that carry populated sample-form "
     "cells",
     "The Stage 1 error was computed across nine substructure families and the "
     "Stage 2 error across the five families that carried populated "
     "sample-form cells",
     "past tense: this sentence describes the superseded comparison", None),
    ("Substructure and sample-form conditioning reduces field-exponent error "
     "16-fold on means and 4.7-fold on medians relative to monolithic "
     "regression.",
     "Substructure conditioning reduces field-exponent error by between one "
     "and about two-fold under leave-one-substructure-out validation, "
     "depending on the cohort, and the interquartile bound reported with "
     "Stage 3 covers the residual in four of four families once the cuprate "
     "families are removed.",
     "same", None),
    # --- two further occurrences found on verification ------------------------
    ("all 110 per-paper anchor groups are identical before and after",
     "all 96 per-paper anchor groups are identical before and after",
     "Sec. III.D, second mention of the anchor count; "
     "phase_3_p31_jc_anchor_per_paper.csv", None),
    ("Conditioning by substructure and sample form reduces the cross-family "
     "field-exponent error by 16-fold on means relative to monolithic "
     "regression on a matched cohort.",
     "Conditioning by substructure reduces the cross-family field-exponent "
     "error by at most about two-fold under leave-one-substructure-out "
     "validation, and by less on some cohorts.",
     "conclusions; analysis/multi_stage_loso.py", None),
    # --- occurrences found on the second verification sweep -------------------
    ("built from 69 papers that contribute fitted critical-current curves, "
     "covering 43 compounds and 4387 extracted data points",
     "built from 62 papers that contribute fitted critical-current curves, "
     "covering 38 compounds and 4146 extracted data points",
     "abstract; provenance_table_fitcohort_full.csv", None),
    ("Sixty-nine papers pass the fittability filters and contribute fitted "
     "curves across 43 compounds and 4387 critical-current data points",
     "Sixty-two papers pass the fittability filters and contribute fitted "
     "curves across 38 compounds and 4146 critical-current data points",
     "Sec. II.A; same", None),
    ("110 per-paper anchors in total, of which 61 fall in the three families "
     "plotted",
     "96 per-paper anchors in total, of which 52 fall in the three families "
     "plotted", "Table II; phase_3_p31_jc_anchor_per_paper.csv", None),
    ("Sample form explains 73%, 60%, and 12% of within-family anchor variance",
     "Sample form explains 77%, 35%, and 12% of within-family anchor variance",
     "Table II; phase_3_p31_variance_decomposition.csv, which gives 0.7687, "
     "0.5988 and 0.1159. Figure 3 already prints 0.77", None),
    ("For iron chalcogenide 11-type materials the ratio is 0.73",
     "For iron chalcogenide 11-type materials the ratio is 0.77",
     "same; the text and the figure disagreed before this change", None),
    ("For iron pnictide 122-type the ratio is 0.60",
     "For iron pnictide 122-type the ratio is 0.35",
     "phase_3_p31_variance_decomposition.csv after the two withdrawals of "
     "2026-09-03; the family loses 7 of its 16 physical samples and the ratio "
     "falls from 0.5988 to 0.3452, close to the 0.3 band boundary", None),
    ("Temperature axis: iron chalcogenide 11-type 0.261 clears the "
     "screening-grade threshold of 1 in the exponent; 122-type 1.092 and "
     "1111-type 1.721 do not; MgB2-class not assessable. Field axis: 0.641, "
     "0.753, 0.973 clear; 1111-type fails.",
     "Temperature axis: iron chalcogenide 11-type 0.588 clears the "
     "screening-grade threshold of 1 in the exponent; 122-type 1.314 and "
     "1111-type 3.120 do not; MgB2-class not assessable. Field axis: "
     "MgB2-class 0.753 and 122-type 0.973 clear; 11-type 1.093 and 1111-type "
     "2.571 do not.",
     "Table II; data/phase_3_p47_compound_leave_out_MAE.csv and "
     "audit/temperature_axis_leave_one_out.csv", None),
    ("which holds 2615 unique PDFs and is the pool from which the 934-article "
     "screened corpus was drawn",
     "which holds 2594 papers and is the pool from which the 934-article "
     "screened corpus was drawn",
     "data/caption_sweep.csv less 21 rows that are not papers", None),
    ("for each of the 2615 unique PDFs in the archive",
     "for each of the 2594 papers in the archive", "same", None),
    ("26 of 80 curves qualify where the pipeline admitted 80, and 13 qualify "
     "once the magnetic-field unit corrections are also applied.",
     "28 of 77 curves qualify where the pipeline admitted 77.",
     "analysis/recompute_supplement_numbers.py. The clause about the unit "
     "corrections is dropped because the pre-correction extraction dataset is "
     "not deposited and the comparison cannot be reproduced", None),
    ("13 for iron chalcogenide 11-type, 16 for iron pnictide 122-type, and 15 "
     "for MgB2-class",
     "10 for iron chalcogenide 11-type, 9 for iron pnictide 122-type, and 15 "
     "for MgB2-class",
     "Fig. 3 caption, markers per family; physical samples per family from "
     "analysis/regenerate_regime_tables.py, which gives 10, 16 and 15 "
     "summing to the 41 markers drawn", None),
]



SUPP_EDITS = [
    ("the 69 source papers that contribute fitted data",
     "the 62 source papers that contribute fitted data",
     "provenance_table_fitcohort_full.csv", None),
    ("The 69 papers listed contribute curves",
     "The 62 papers listed contribute curves", "same", None),
    ("The median ratio of measured maximum to assigned scale is 0.86, and for "
     "15 of 77 curves it exceeds 0.9.",
     "The median ratio of measured maximum to assigned scale is 0.80, and for "
     "15 of 94 curves it exceeds 0.9.",
     "analysis/recompute_supplement_numbers.py", None),
    ("Thirty-one of the 80 assigned scales are an irreversibility field or are "
     "unlabelled rather than a confirmed upper critical field.",
     "Thirty of the 77 assigned scales are an irreversibility field or are "
     "unlabelled rather than a confirmed upper critical field.", "same", None),
    ("26 of 80 curves qualify where the pipeline admitted 80",
     "28 of 77 curves qualify where the pipeline admitted 77", "same", None),
    ("phase_3_p31_jc_anchor_per_paper.csv, the 110-row file behind the "
     "variance-decomposition diagnostic",
     "phase_3_p31_jc_anchor_per_paper.csv, the 96-row file behind the "
     "variance-decomposition diagnostic",
     "phase_3_p31_jc_anchor_per_paper.csv", None),
    ("The per-paper leave-one-out Stage 2 validation gives a post-expansion "
     "mean absolute error of 0.994 in \u03b2H, with a 95% bootstrap confidence "
     "interval of [0.495, 1.661]. The bootstrap uses 5000 iterations with seed "
     "42, applying the percentile method throughout. The validation cohort "
     "contains 99 fits across five substructures and 17 papers. It comprises "
     "95 field-axis partial fits that pass the physicality checks, plus four "
     "additional FeSe pure-compound fits incorporated after an expansion of "
     "the applicability criterion. Twelve additional field-axis fits are "
     "excluded because they fail the field-axis applicability boundary. The "
     "corresponding pre-expansion estimate was 1.069, with a 95% confidence "
     "interval of [0.632, 1.816]. The post-expansion result therefore lowers "
     "the point estimate by 0.075 and lowers the upper bound by 0.155; the "
     "interval width is essentially unchanged, 1.166 against 1.184. The "
     "overlapping intervals indicate cohort-stable performance after the "
     "expansion.",
     "The per-paper leave-one-out validation is regenerated by "
     "analysis/compound_leave_one_out.py and deposited as "
     "audit/per_paper_field_validation.csv. We report both predictors, because "
     "they no longer favour the conditioned one. The Stage 2 predictor, "
     "conditioned on substructure and sample form, gives a mean absolute error "
     "of 1.257 in \u03b2H over 82 scored fits from 16 papers, with a 95% "
     "bootstrap confidence interval of [1.006, 1.612]. A pooled median with no "
     "conditioning gives 1.158 over 94 fits from the same 16 papers, with "
     "[0.900, 1.497]. The bootstrap uses 5000 iterations with seed 20260901 "
     "and the percentile method throughout. The two are not scored on the same "
     "cohort, since the conditioned predictor cannot score a fit whose "
     "sample-form cell is unpopulated, so this is not a like-for-like "
     "comparison; but the conditioned predictor no longer has the lower error "
     "and we state that rather than reporting only the figure that flatters "
     "it. An earlier version of this section quoted 0.994 against a "
     "pre-expansion 1.069 on a 99-fit cohort. Neither is reproducible from the "
     "deposited data and both are withdrawn.",
     "audit/per_paper_field_validation.csv", None),
    ("The cuprate substructures account for 62% of the residual mass but "
     "represent only 44% of the substructure cohort. Three cuprate "
     "substructures hit or approach the imposed regression ceiling: RBCO "
     "(REBa2Cu3O7, rare-earth\u2013barium\u2013copper\u2013oxide) and BSCCO "
     "(Bi\u2013Sr\u2013Ca\u2013Cu\u2013O) at \u03b2 = 30, HBCCO "
     "(Hg\u2013Ba\u2013Ca\u2013Cu\u2013O) at \u03b2 = 25.81; LSCO "
     "(La2\u2212xSrxCuO4, lanthanum\u2013strontium\u2013copper\u2013oxide) is "
     "physical at \u03b2 = 5.50.",
     "The cuprate substructures account for 59% of the residual mass but "
     "represent only 43% of the substructure cohort. RBCO (REBa2Cu3O7, "
     "rare-earth\u2013barium\u2013copper\u2013oxide) sits at the imposed "
     "regression ceiling of \u03b2 = 30 on all 3 of its fits and BSCCO "
     "(Bi\u2013Sr\u2013Ca\u2013Cu\u2013O) on all 20 of its fits, while LSCO "
     "(La2\u2212xSrxCuO4, lanthanum\u2013strontium\u2013copper\u2013oxide) is "
     "physical. The ceiling is not confined to the cuprates: 4 of 46 iron "
     "chalcogenide 11-type fits and 1 of 22 iron pnictide 1111-type fits also "
     "reach it.",
     "audit/multi_stage_loso.csv for the residual shares; "
     "phase_3_form3_fits_partial_cohortB_v2.csv for the ceiling counts. HBCCO "
     "left the cohort with an earlier withdrawal", None),
    ("A caption-scoped screen over the 2615 unique PDFs in the archive",
     "A caption-scoped screen over the 2594 papers in the archive", "same", None),
    ("a compound leave-one-out mean absolute error of 5.13 in \u03b2H under the "
     "expanded compound-diversity validation, and 3.07 at the cohort scope of "
     "the main-text Stage 2 table",
     "a compound leave-one-out mean absolute error of 2.571 in \u03b2H at the "
     "cohort scope of the main-text Stage 2 table",
     "data/phase_3_p47_compound_leave_out_MAE.csv. The expanded "
     "compound-diversity figure is dropped: its cohort is not deposited", None),
]




# Supplement Table S1, rebuilt from provenance_table_fitcohort_full.csv by
# analysis/rebuild_supplement_table_s1.py. The table was static and still
# described a 69-paper cohort with three families that no longer contribute.
TABLE_S1 = [
    ["Iron chalcogenide 11", "17", "7", "2", "8", "0", "10", "5", "980"],
    ["Iron pnictide 111", "2", "2", "0", "0", "0", "1", "0", "85"],
    ["Iron pnictide 1111", "13", "6", "0", "7", "0", "10", "6", "1149"],
    ["Iron pnictide 122", "21", "14", "0", "7", "0", "10", "7", "1444"],
    ["Conventional AlB2", "4", "3", "0", "1", "0", "2", "4", "267"],
    ["Cuprate BSCCO", "3", "0", "0", "3", "0", "3", "0", "152"],
    ["Cuprate LSCO", "1", "0", "0", "1", "0", "1", "0", "48"],
    ["Cuprate RBCO", "1", "0", "0", "1", "0", "1", "0", "21"],
    ["TOTAL", "62", "32", "2", "28", "0", "38", "22", "4146"],
]


def set_cell(cell, value):
    """Write a value into a table cell, keeping the first run's formatting."""
    par = cell.paragraphs[0]
    if not par.runs:
        par.add_run(value)
        return
    par.runs[0].text = value
    for r in par.runs[1:]:
        r.text = ""
    for extra in cell.paragraphs[1:]:
        for r in extra.runs:
            r.text = ""


def rebuild_table_s1(doc, report):
    """Replace the provenance table's data rows, and drop the surplus rows."""
    target = None
    for tb in doc.tables:
        if tb.rows and "Substructure family" in tb.rows[0].cells[0].text:
            target = tb
            break
    if target is None:
        report.append(("supplement", False, "Table S1 (provenance)", "", "not found"))
        return ["Table S1"]
    data = target.rows[1:]
    for i, vals in enumerate(TABLE_S1):
        if i >= len(data):
            report.append(("supplement", False, "Table S1 row %d" % i, "", "no row to write into"))
            return ["Table S1"]
        for j, v in enumerate(vals):
            if j < len(data[i].cells):
                set_cell(data[i].cells[j], v)
    for row in data[len(TABLE_S1):]:
        row._element.getparent().remove(row._element)
    report.append(("supplement", True, "Table S1 (provenance), %d rows" % len(TABLE_S1),
                   "TOTAL 62 papers, 38 compounds, 4146 points",
                   "analysis/rebuild_supplement_table_s1.py"))
    return []


def replace_in_paragraph(par, find, repl):
    """Replace one occurrence across runs, keeping formatting outside the span."""
    runs = par.runs
    if not runs:
        return False
    text = "".join(r.text for r in runs)
    at = text.find(find)
    if at < 0:
        return False
    end = at + len(find)
    pos, first = 0, None
    for i, r in enumerate(runs):
        s, e = pos, pos + len(r.text)
        if first is None and e > at:
            first = i
            head = r.text[:at - s]
            tail = r.text[end - s:] if e >= end else ""
            r.text = head + repl + tail
            if e >= end:
                return True
        elif first is not None:
            if e <= end:
                r.text = ""
            else:
                r.text = r.text[end - s:]
                return True
        pos = e
    return True


def apply(doc, edits, label, report):
    misses = []
    for find, repl, why, loc in edits:
        done = False
        if loc and loc.startswith("table:"):
            _, ti, ri, ci = loc.split(":")
            cell = doc.tables[int(ti)].rows[int(ri)].cells[int(ci)]
            for par in cell.paragraphs:
                if replace_in_paragraph(par, find, repl):
                    done = True
                    break
        else:
            for par in doc.paragraphs:
                if replace_in_paragraph(par, find, repl):
                    done = True
                    break
            if not done:
                for tb in doc.tables:
                    for row in tb.rows:
                        for cell in row.cells:
                            for par in cell.paragraphs:
                                if replace_in_paragraph(par, find, repl):
                                    done = True
                                    break
        report.append((label, done, find[:72], repl[:72], why))
        if not done:
            misses.append(find)
    return misses


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--ms", required=True)
    ap.add_argument("--supp", required=True)
    ap.add_argument("--out-dir", default=".")
    ap.add_argument("--dry-run", action="store_true")
    args = ap.parse_args()

    report, misses = [], []
    ms = docx.Document(args.ms)
    misses += apply(ms, MS_EDITS, "manuscript", report)
    supp = docx.Document(args.supp)
    misses += apply(supp, SUPP_EDITS, "supplement", report)
    misses += rebuild_table_s1(supp, report)

    for lab, ok, f, r, why in report:
        print("%-11s %-7s %s\n%-11s %-7s -> %s\n            %s\n"
              % (lab, "applied" if ok else "MISSED", f, "", "", r, why))

    if misses:
        print("%d edit(s) found no target. Nothing written." % len(misses))
        for m in misses:
            print("   " + m[:100])
        return 1

    if args.dry_run:
        print("%d edits would be applied. Nothing written." % len(report))
        return 0

    for src, doc in ((args.ms, ms), (args.supp, supp)):
        base = os.path.basename(src)
        stem, ext = os.path.splitext(base)
        out = os.path.join(args.out_dir, stem + "_corrected" + ext)
        keep = os.path.join(args.out_dir, stem + "_asfound" + ext)
        if not os.path.exists(keep):
            shutil.copy2(src, keep)
        doc.save(out)
        print("written %s  (original preserved as %s)" % (out, keep))
    return 0


if __name__ == "__main__":
    sys.exit(main())
