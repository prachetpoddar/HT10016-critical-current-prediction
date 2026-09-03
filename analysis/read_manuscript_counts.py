#!/usr/bin/env python3
"""Read the counts the manuscript actually prints, and record where each came from.

Why this exists. `verify_deposit.py` carried a MANUSCRIPT dictionary of counts
described as "what the manuscript prints". It was not. Every value in it equals
what the deposit itself held on the day the dictionary was written, so the check
compared the deposit with a snapshot of the deposit and would pass while the
manuscript said something else entirely. Table I of HT10016_revised.docx prints
69, 43, 4387, 419, 95 and 110 where that dictionary held 65, 40, 4247, 414, 88
and 105.

A number asserted against the manuscript has to be read out of the manuscript.
This script does that, writes the value with the file, the table or paragraph it
came from, and the date, and `verify_deposit.py` reads the result.

    python analysis/read_manuscript_counts.py --docx HT10016_revised.docx \
        --out audit/manuscript_printed_counts.csv
"""
import argparse
import csv
import datetime
import os
import re

import docx

# Table I row label -> the key verify_deposit uses.
TABLE_I = {
    "Papers contributing fitted curves": "fitted_curve_papers",
    "Distinct compounds with fitted curves": "fitted_curve_compounds",
    "Critical-current data points extracted": "extracted_points",
    "Temperature-axis partial fits": "temperature_axis_fits",
    "Field-axis partial fits passing physicality": "field_axis_fits_ok",
    "Per-paper anchors behind Fig. 3": "anchor_rows",
    "Candidate compounds evaluated": "candidate_compounds",
}


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--docx", required=True)
    ap.add_argument("--out", required=True)
    args = ap.parse_args()

    d = docx.Document(args.docx)
    src = os.path.basename(args.docx)
    stamp = datetime.date.today().isoformat()
    rows = []

    for r in d.tables[0].rows[1:]:
        cells = [c.text.strip() for c in r.cells]
        label = re.sub(r"\s*\(.*?\)\s*$", "", cells[0]).strip()
        key = TABLE_I.get(label)
        if not key:
            continue
        m = re.search(r"[0-9][0-9,]*", cells[1])
        if not m:
            continue
        rows.append(dict(quantity=key, value=int(m.group(0).replace(",", "")),
                         printed_as=cells[1].strip(),
                         located_in="Table I, row '%s'" % cells[0].strip(),
                         document=src, read_on=stamp))

    # A count that appears only in prose. The field-axis source-paper count is
    # printed in the "What it supports" column of Table I rather than its own row.
    for r in d.tables[0].rows[1:]:
        cells = [c.text.strip() for c in r.cells]
        m = re.search(r"from (\d+) source papers", cells[2] if len(cells) > 2 else "")
        if m:
            rows.append(dict(quantity="field_axis_ok_papers", value=int(m.group(1)),
                             printed_as=m.group(0),
                             located_in="Table I, 'What it supports' column of "
                                        "row '%s'" % cells[0].strip(),
                             document=src, read_on=stamp))

    with open(args.out, "w", newline="") as fh:
        w = csv.DictWriter(fh, fieldnames=["quantity", "value", "printed_as",
                                           "located_in", "document", "read_on"],
                           lineterminator="\n")
        w.writeheader()
        w.writerows(rows)

    for r in rows:
        print("%-28s %-8s  %s" % (r["quantity"], r["value"], r["located_in"]))
    print("\nwritten to %s" % args.out)


if __name__ == "__main__":
    main()
