"""
apply_dispatch_count_edits.py

The candidate-dispatch accounting, corrected across all three artifacts.

Every value here moved because of the reduced-field gate and the record
withdrawals, and every one of them is a count or a share that the deposit
computes. Each replacement is built from the deposited tables at run time and
asserted before it is written; the script refuses to write if any edit misses.

The provenance of the two figures that change most:

  * 239 records / 185 compounds / 2151 tuples was the cohort before the
    withdrawals. It is 233 / 183 / 2097 now, and 233 x 9 = 2097 holds as it
    did before.
  * 0.39 dex was the median 95% interval width over a pre-gate population of
    1671 rows carrying either an emitted or a withheld value. Over the 256
    rows the gate actually emits it is 0.825 dex, a factor of about 6.7 rather
    than 2.5. analysis/manuscript_figure_5.py already computes the envelope
    half-width from the deposit and draws 0.41; only the captions still said
    0.19 and 0.39.

Held back, because they are claims and not counts. Two of the three dispatch
families now emit nothing at all: every one of iron chalcogenide 11-type's 441
grid points is refused (248 reduced field, 153 above Tc, 40 family field-axis)
and all 711 of iron pnictide 122-type's (378 Hc2 unavailable, 294 reduced
field, 39 above Tc). The emitted grid carries one field, 5 T. So these passages
describe dispatches that no longer happen, and rewriting them is a decision
about what the paper claims:

  * "Of the 185, 125 receive at least one non-refused prediction target ...
    85 MgB2-class, 31 iron chalcogenide 11-type, and 9 iron pnictide 122-type"
  * "123 rather than 125 of the 185 candidates carry a prediction we are
    willing to report"
  * "a refusal decision for each of 185 candidate compounds ... of which 125
    compounds receive at least one dispatched target"
  * "At 4.2 K and 0.1 T, the 87 dispatched MgB2-class prediction records ..."
    The 87 records and 85 compounds are right; the field is 5 T, and the
    chalcogenide and 122 counts in the same sentence are now zero.
  * "Family-level screening outcome. We report family medians at 4.2 K and the
    lowest evaluated field of 0.1 T ... 6.00, 5.75, 5.32". Nothing is
    dispatched at 0.1 T. The one surviving family median is 4.98 for the MgB2
    class at 4.2 K and 5 T, over 87 records.
  * "For iron chalcogenide 11-type, sample-form conditioning is mandatory, so
    the dispatch uses the single-crystal empirical cell ..."
  * The supplement's "On the compound basis, 125 of the 185 receive at least
    one non-refused prediction target"

Usage:
    python3 analysis/apply_dispatch_count_edits.py \\
        --ms MS.docx --supp SUPP.docx --letter LETTER.docx \\
        --letter-md LETTER.md --out-dir DIR
"""
import argparse
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import docx  # noqa: E402
import pandas as pd  # noqa: E402

from apply_manuscript_edits import apply  # noqa: E402


def numbers():
    t = pd.read_csv(os.path.join("data",
                                 "phase_3_p56_candidate_tier_assignment.csv"))
    p = pd.read_csv(os.path.join("data",
                                 "phase_3_p57_de_novo_predictions.csv"),
                    low_memory=False)
    em = p[p.refusal_flag.fillna("") == ""]
    width = float((em.predicted_log_Jc_upper_95
                   - em.predicted_log_Jc_lower_95).median())
    has = int(t.Hc2_T.notna().sum())
    return dict(
        records=len(t),
        compounds=int(p.compound_formula.nunique()),
        tuples=len(p),
        grid=len(p) // len(t),
        refused_cal=int((t.tier == "refused_calibration_domain").sum()),
        retained=int((t.tier != "refused_calibration_domain").sum()),
        below=int((t.tier == "graded_confidence").sum()),
        at_or_above=int((t.tier == "high_confidence").sum()),
        dispatched=int(em.compound_formula.nunique()),
        hc2_records=has,
        hc2_share=100.0 * has / len(t),
        width=width,
        half=width / 2,
        factor=10 ** width,
    )


def edits(N):
    return dict(
        manuscript=[
            ("The candidate set contains 239 candidate records covering 185 "
             "distinct compounds",
             "The candidate set contains %d candidate records covering %d "
             "distinct compounds" % (N["records"], N["compounds"]),
             "candidate cohort, recomputed", None),
            ("it returned no match for any of the 239 candidate records",
             "it returned no match for any of the %d candidate records"
             % N["records"], "same record count", None),
            ("21 of the 239 candidate records carry a transition-temperature "
             "anchor below 4.2 K",
             "%d of the %d candidate records carry a transition-temperature "
             "anchor below 4.2 K" % (N["refused_cal"], N["records"]),
             "same record count", None),
            ("the framework ranks 185 candidate compounds",
             "the framework ranks %d candidate compounds" % N["compounds"],
             "same compound count", None),
            ("an explicit refusal decision for each of 185 candidate "
             "compounds and each prediction target, of which 125 compounds "
             "receive at least one dispatched target",
             "an explicit refusal decision for each of %d candidate compounds "
             "and each prediction target, of which %d compounds receive at "
             "least one dispatched target" % (N["compounds"], N["dispatched"]),
             "conclusion, recomputed", None),
            ("is 0.39 dex in log10 Jc, a factor of about 2.5 in Jc",
             "is %.2f dex in log10 Jc, a factor of about %.1f in Jc"
             % (N["width"], N["factor"]),
             "median interval width over the emitted rows, not the pre-gate "
             "population", None),
            ("Shaded envelopes are drawn at a constant half-width of 0.19 dex "
             "around each curve, which is half the median full width of the "
             "95% bootstrap confidence interval quoted as 0.39 dex in the text",
             "Shaded envelopes are drawn at a constant half-width of %.2f dex "
             "around each curve, which is half the median full width of the "
             "95%% bootstrap confidence interval quoted as %.2f dex in the text"
             % (N["half"], N["width"]),
             "Fig. 5 already draws the recomputed half-width; only the caption "
             "still said 0.19", None),
        ],
        supplement=[
            ("Across the three dispatched families, 179 of the 239 candidate "
             "records have upper-critical-field coverage, corresponding to "
             "74.9% coverage on the record basis",
             "Across the three dispatched families, %d of the %d candidate "
             "records have upper-critical-field coverage, corresponding to "
             "%.1f%% coverage on the record basis"
             % (N["hc2_records"], N["records"], N["hc2_share"]),
             "Hc2 coverage on the record basis, recomputed", None),
            ("The complete prediction table includes 239 candidate records "
             "covering 185 distinct compounds, evaluated at nine target grid "
             "points, giving 2151 candidate-grid tuples",
             "The complete prediction table includes %d candidate records "
             "covering %d distinct compounds, evaluated at %s target grid "
             "points, giving %d candidate-grid tuples"
             % (N["records"], N["compounds"],
                "nine" if N["grid"] == 9 else str(N["grid"]), N["tuples"]),
             "prediction table size, recomputed", None),
            ("Of the 239 records, 21 fail this test and 218 are retained",
             "Of the %d records, %d fail this test and %d are retained"
             % (N["records"], N["refused_cal"], N["retained"]),
             "calibration-domain screen, recomputed", None),
            ("130 records fall below their family range and 88 fall at or "
             "above it, of which 87 are inside the range and one above it",
             "%d records fall below their family range and %d fall at or "
             "above it, of which %d are inside the range and one above it"
             % (N["below"], N["at_or_above"], N["at_or_above"] - 1),
             "calibration-range classification, recomputed", None),
            ("Missing upper-critical-field anchors produce refusals for 25.1% "
             "of candidate-grid tuples. Target temperatures above the "
             "transition temperature produce refusals for 10.5% of "
             "candidate-grid tuples.",
             "Predictions below the validated reduced field produce refusals "
             "for %.1f%% of candidate-grid tuples, missing "
             "upper-critical-field anchors for %.1f%%, and target temperatures "
             "above the transition temperature for %.1f%%."
             % (_share(), _share("Hc2_unavailable"), _share("T_above_Tc")),
             "four codes fire; the reduced-field code is the largest and was "
             "unmentioned", None),
            ("much smaller than the median bootstrap confidence-interval "
             "width of 0.39 dex, a factor of about 2.5 in Jc",
             "much smaller than the median bootstrap confidence-interval "
             "width of %.2f dex, a factor of about %.1f in Jc"
             % (N["width"], N["factor"]),
             "same interval width", None),
            ("0.30 dex is a factor of 2, and 0.39 dex is a factor of about 2.5",
             "0.30 dex is a factor of 2, and %.2f dex is a factor of about %.1f"
             % (N["width"], N["factor"]),
             "glossary example, kept equal to the quoted width", None),
        ],
        letter=[
            ("The candidate cohort comprises 185 distinct compounds, and the "
             "previously reported 239 and 218 are record counts in which one "
             "compound may appear more than once.",
             "The candidate cohort comprises %d distinct compounds, and the "
             "%d evaluated records and %d retained after the calibration "
             "screen are record counts in which one compound may appear more "
             "than once." % (N["compounds"], N["records"], N["retained"]),
             "candidate cohort, recomputed", None),
            ("could be read as a ranking of 185 candidate compounds",
             "could be read as a ranking of %d candidate compounds"
             % N["compounds"], "same compound count", None),
        ],
    )


_P57 = None


def _share(flag="H_below_validated_reduced_field"):
    global _P57
    if _P57 is None:
        _P57 = pd.read_csv(os.path.join(
            "data", "phase_3_p57_de_novo_predictions.csv"), low_memory=False)
    f = _P57.refusal_flag.fillna("")
    return 100.0 * (f == flag).sum() / len(_P57)


def _rewrap(find):
    return re.compile(r"\s+".join(re.escape(w) for w in find.split()))


def _flat(doc):
    """All the text of a python-docx document, whitespace normalised."""
    parts = [par.text for par in doc.paragraphs]
    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                parts.extend(par.text for par in cell.paragraphs)
    return re.sub(r"\s+", " ", " ".join(parts))


def _already(text, repl):
    """True when the replacement is already in place, so a miss is idempotence
    and not a failure. Run twice, this script must report done and not refuse."""
    return _rewrap(repl).search(text) is not None


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--ms", required=True)
    ap.add_argument("--supp", required=True)
    ap.add_argument("--letter", required=True)
    ap.add_argument("--letter-md", required=True)
    ap.add_argument("--out-dir", default=".")
    ap.add_argument("--dry-run", action="store_true")
    a = ap.parse_args()
    if not os.path.isdir("data"):
        sys.exit("run from the repository root")

    N = numbers()
    print("recomputed from the deposit\n")
    for k in ("records", "compounds", "tuples", "retained", "below",
              "at_or_above", "hc2_records"):
        print("   %-16s %8d" % (k, N[k]))
    print("   %-16s %8.1f%%" % ("hc2_share", N["hc2_share"]))
    print("   %-16s %8.3f dex  (factor %.1f)" % ("interval width", N["width"],
                                                 N["factor"]))
    print()

    E = edits(N)
    report, misses = [], []
    docs, applied = {}, 0
    for label, path in (("manuscript", a.ms), ("supplement", a.supp),
                        ("letter", a.letter)):
        d = docx.Document(path)
        sub = []
        raw = apply(d, E[label], label, sub)
        text = _flat(d)
        for entry in sub:
            lb, done, find, repl, why = entry
            if not done:
                full = next(r for f, r, _, _ in E[label] if f[:72] == find)
                if _already(text, full):
                    report.append((lb, "done", find, repl, why))
                    continue
                misses.append(find)
                report.append((lb, "MISS", find, repl, why))
            else:
                applied += 1
                report.append((lb, "ok", find, repl, why))
        docs[label] = d

    md = open(a.letter_md, encoding="utf-8").read()
    for find, repl, why, _ in E["letter"]:
        pat = _rewrap(find)
        if pat.search(md):
            md = pat.sub(repl.replace("\\", "\\\\"), md, count=1)
            applied += 1
            report.append(("letter, markdown", "ok", find[:66], repl[:66], why))
        elif _already(re.sub(r"\s+", " ", md), repl):
            report.append(("letter, markdown", "done", find[:66], repl[:66],
                            why))
        else:
            report.append(("letter, markdown", "MISS", find[:66], repl[:66],
                            why))
            misses.append(find)

    for label, state, find, repl, why in report:
        if state == "done":
            print("   %-18s done %s   (already applied)" % (label, find))
        else:
            print("   %-18s %-4s %s\n        -> %s"
                  % (label, state, find, repl))
    print("\n   %d edit(s) applied this run" % applied)
    print()
    if misses:
        print("%d edit(s) missed; nothing written" % len(misses))
        return 1
    if a.dry_run:
        print("dry run, nothing written")
        return 0
    out = {"manuscript": os.path.basename(a.ms),
           "supplement": os.path.basename(a.supp),
           "letter": os.path.basename(a.letter)}
    for label, d in docs.items():
        d.save(os.path.join(a.out_dir, out[label]))
        print("   wrote %s" % out[label])
    with open(os.path.join(a.out_dir, os.path.basename(a.letter_md)), "w",
              encoding="utf-8") as fh:
        fh.write(md)
    print("   wrote %s" % os.path.basename(a.letter_md))
    return 0


if __name__ == "__main__":
    sys.exit(main())
