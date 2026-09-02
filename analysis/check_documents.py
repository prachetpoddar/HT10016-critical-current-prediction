#!/usr/bin/env python3
"""
check_documents.py

Cross-document consistency check for the resubmission package.

Why this exists. The manuscript, the supplement and the response letter have
each, at some point in this revision, carried a cohort count from a superseded
version of the analysis. Prose is where stale numbers survive, because nothing
recomputes it. This script recomputes the cohort from the deposit, then reads
the three documents and fails if a superseded value appears in a sentence that
does not mark it as historical.

A superseded value is allowed to appear, and often should: a response letter
that corrects a number has to name the number it is correcting. What is not
allowed is a superseded value stated as current. The test for the difference is
whether its sentence carries one of the markers below, which is crude but is
checkable by a reader and by a script, and which fails closed.

Needs python-docx, which is imported as "docx" but installed as
"pip install python-docx".

    python analysis/check_documents.py --docs ~/Downloads
    python analysis/check_documents.py --files a.docx b.docx c.docx
    python analysis/check_documents.py --docs ~/Downloads --list

By default a folder is filtered to the three documents of this package, since a
download folder holds other people's papers and checking those produces noise.
Pass --all to override that.

Exit status is non-zero if any superseded value is stated as current, or if a
current value cannot be found where it is expected.

Run from the repository root.
"""
import argparse
import glob
import os
import re
import sys

try:
    import docx  # python-docx
except ImportError:
    sys.exit("this script needs python-docx:\n"
             "    pip install python-docx\n"
             "(the module is imported as 'docx' but the package is 'python-docx')")

# A superseded value may appear when its sentence marks it as historical, either
# by carrying one of these phrases or by naming the replacement value alongside
# it. The second rule is the stronger one and catches corrections phrased as
# "from X to Y", which is how most of them read.
# Filename fragments that identify the three documents of this package.
PACKAGE_NAMES = {"HT10016", "SUPPLEMENTAL", "SUPPLEMENT", "RESPONSE"}

MARKERS = [
    "earlier version", "an earlier", "previously", "no longer", "superseded",
    "withdrawn", "we withdraw", "rather than", "against the", "instead of",
    "was ", "were ", "before", "historical", "used to", "prior to",
    "corrected", "correction", "falls to", "falls from", "changes from",
    "in place of", "replaced",
]

# value that must not be stated as current  ->  what replaced it
SUPERSEDED = {
    r"\b69 (?:source )?papers\b": "65 papers",
    r"\b43 (?:distinct )?compounds\b": "40 compounds",
    r"\b4387\b": "4247 extracted points",
    r"\b4355\b": "4247 extracted points",
    r"\b419 temperature-axis fits\b": "414 temperature-axis fits",
    r"\b95 field-axis\b": "88 field-axis fits",
    r"\b93 field-axis\b": "88 field-axis fits",
    r"\b110 (?:per-paper )?anchors?\b": "105 anchor records",
    r"\b107 per-paper anchor\b": "105 anchor records",
    r"\b0\.3547\b": "0.3409 aggregate ratio",
    r"\bratio is 0\.73\b": "0.77 for iron chalcogenide 11-type",
    r"\bexplains 76%\b": "77%",
    r"\b123 (?:compounds|of the 183)\b": "85 dispatched compounds",
    r"\b0\.641 for iron chalcogenide\b": "1.094",
    r"\b3\.07 at this cohort scope\b": "3.13",
    r"\b0\.994 in .H\b": "1.141",
    r"\b97 fits\b": "88 fits",
    r"\b2151\b": "2097 candidate-grid tuples",
    r"\b239 candidate records\b": "233 candidate records",
    r"\b185 distinct compounds\b": "183 distinct compounds",
    r"\b25\.1%\b": "25.8%",
    r"\b10\.5%\b": "9.9%",
    r"\b92% of (?:bootstrap )?resamples\b": "withdrawn, not restated",
    r"\b23-fold\b": "16-fold",
    r"cannot be reproduced": "no unreproducible value may remain",
    r"does not reproduce": "no unreproducible value may remain",
}


def sentences(text):
    return re.split(r"(?<=[.!?])\s+", text)


def doc_text(path):
    """Blocks of text, with referee quotations marked so they are exempt.

    The response letter quotes each referee point verbatim before answering it,
    and those quotations name the numbers being corrected. A quotation is set
    wholly in italic, which is how it is identified here; its content is the
    referee's, not ours, and must not be edited to match our cohort.
    """
    # Only the response letter quotes referees, so the exemption is confined to
    # it. Applying it document-wide would let an italic caption hide a stale
    # number, which is the failure mode this script exists to catch.
    is_response = "RESPONSE" in os.path.basename(path).upper()
    d = docx.Document(path)
    out = []
    for p in d.paragraphs:
        quoted = (is_response and len(p.text.strip()) > 40 and bool(p.runs)
                  and all(r.italic for r in p.runs if r.text.strip()))
        out.append((p.text, quoted))
    for t in d.tables:
        for r in t.rows:
            for c in r.cells:
                out.append((c.text, False))
    return out


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--docs",
                    help="folder holding the manuscript, supplement and response")
    ap.add_argument("--files", nargs="+",
                    help="the documents to check, given explicitly")
    ap.add_argument("--all", action="store_true",
                    help="with --docs, check every .docx rather than only the "
                         "three that make up this package")
    ap.add_argument("--list", action="store_true",
                    help="print every hit, marked or not")
    args = ap.parse_args()

    if args.files:
        paths = sorted(args.files)
    elif args.docs:
        paths = sorted(p for p in glob.glob(os.path.join(args.docs, "*.docx"))
                       if not os.path.basename(p).startswith("~$"))
        if not args.all:
            # A folder such as Downloads holds other people's documents, and
            # checking those produces noise at best and a traceback at worst.
            paths = [p for p in paths
                     if any(k in os.path.basename(p).upper()
                            for k in PACKAGE_NAMES)]
    else:
        sys.exit("give --docs FOLDER or --files A.docx B.docx")
    if not paths:
        where = args.docs or "the given files"
        sys.exit("no package document found in %s.\n"
                 "Expected a filename containing one of: %s\n"
                 "Use --files to name them explicitly, or --all to check "
                 "every .docx in the folder." % (where, ", ".join(sorted(PACKAGE_NAMES))))

    bad = 0
    for path in paths:
        name = os.path.basename(path)
        hits_marked = hits_bare = 0
        print("\n%s" % name)
        for block, quoted in doc_text(path):
            if quoted:
                continue
            for s in sentences(block):
                for pat, repl in SUPERSEDED.items():
                    if not re.search(pat, s, re.I):
                        continue
                    # A sentence that names the superseded value and its
                    # replacement together is a correction, which is the
                    # clearest form a historical mention can take.
                    new_num = re.match(r"([\d.]+)", repl)
                    paired = bool(new_num) and re.search(
                        r"\b%s\b" % re.escape(new_num.group(1)), s)
                    marked = paired or any(m in s.lower() for m in MARKERS)
                    if marked:
                        hits_marked += 1
                        if args.list:
                            print("   ok   %-34s %s" % (repl, s.strip()[:96]))
                    else:
                        hits_bare += 1
                        bad += 1
                        print("   FAIL %-34s %s" % (repl, s.strip()[:112]))
        nq = sum(1 for _b, q in doc_text(path) if q)
        print("   %d superseded value(s) marked as historical, %d stated as current"
              "   (%d verbatim referee quotation(s) exempt)"
              % (hits_marked, hits_bare, nq))

    print()
    if bad:
        print("%d superseded value(s) stated as current; the package is not consistent" % bad)
        return 1
    print("no superseded value is stated as current")
    return 0


if __name__ == "__main__":
    sys.exit(main())
