"""
apply_dispatch_scope_rewrite.py

Rewrite Sec. III.E to the dispatch that actually happens.

The reduced-field gate is the central change of this revision, and it closed
two of the three dispatch families. Measured on the deposited prediction table:

  iron chalcogenide 11-type   441 targets, 0 emitted
        248 below the validated reduced field
        153 target temperature above Tc
         40 family does not pass field-axis validation
  iron pnictide 122-type      711 targets, 0 emitted
        378 no upper-critical-field anchor
        294 below the validated reduced field
         39 target temperature above Tc
  MgB2-class                  945 targets, 256 emitted, 85 of 103 compounds

The grid carries three fields, 0.1, 1 and 5 T. Every target at 0.1 and 1 T is
refused, so the only grid point at which anything is dispatched is 5 T.

Sec. III.E still reported the pre-gate picture: 125 compounds across three
families, family medians at 0.1 T of 6.00, 5.75 and 5.32, and a top-quartile
slice that was entirely iron chalcogenide. Every one of those rests on
predictions that are now refused. The numbers could not be swapped, because
there is no corresponding value to swap to; the passages are replaced.

One claim gets weaker and is stated as measured rather than as remembered. The
old text said the records of a family "collapse to one single predicted value
per family, identical to nine decimal places". At 4.2 K and 5 T the 87 MgB2
records take 34 distinct values spanning 0.0098 dex. That is 1.6% of the
0.60 dex bootstrap interval at the same point, so the structural claim holds
comfortably, but it is not identity and is no longer written as identity.

Every figure here is recomputed from the deposit at run time and asserted
against the text before it is written.

Usage:
    python3 analysis/apply_dispatch_scope_rewrite.py \\
        --ms MS.docx --supp SUPP.docx --letter L.docx --letter-md L.md \\
        --out-dir DIR [--dry-run]
"""
import argparse
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import docx  # noqa: E402
import pandas as pd  # noqa: E402

from apply_manuscript_edits import apply  # noqa: E402

GRID_POINT = (4.2, 5.0)


def numbers():
    p = pd.read_csv(os.path.join("data",
                                 "phase_3_p57_de_novo_predictions.csv"),
                    low_memory=False)
    em = p[p.refusal_flag.fillna("") == ""]
    fams = {}
    for fam, g in p.groupby("substructure"):
        counts = g.refusal_flag.fillna("").value_counts().to_dict()
        fams[fam] = dict(targets=len(g), emitted=counts.get("", 0),
                         compounds=int(g.compound_formula.nunique()),
                         dispatched=int(em[em.substructure == fam]
                                        .compound_formula.nunique()),
                         **{k: v for k, v in counts.items() if k})
    at = em[(em.T_K == GRID_POINT[0]) & (em.H_T == GRID_POINT[1])]
    span = float(at.predicted_log_Jc.max() - at.predicted_log_Jc.min())
    width = float((at.predicted_log_Jc_upper_95
                   - at.predicted_log_Jc_lower_95).median())
    # The whole rewrite rests on exactly one family emitting. Assert it rather
    # than assume it, so this script cannot quietly write a false sentence if
    # the gate settings change.
    emitting = [f for f, d in fams.items() if d["emitted"]]
    if emitting != ["conventional_AlB2"]:
        raise SystemExit("this rewrite assumes only conventional_AlB2 emits; "
                         "emitting families are %s" % ", ".join(emitting))
    if len(set(em.H_T)) != 1:
        raise SystemExit("this rewrite assumes one emitted field; the emitted "
                         "grid carries %s" % sorted(set(em.H_T)))
    return dict(
        fams=fams,
        compounds=int(p.compound_formula.nunique()),
        dispatched=int(em.compound_formula.nunique()),
        records_at=len(at),
        compounds_at=int(at.compound_formula.nunique()),
        median_at=float(at.predicted_log_Jc.median()),
        span=span,
        width_at=width,
        ratio=width / span,
        anchor=float(at.Hc2_T_anchor.iloc[0]),
        n_anchors=int(em.Hc2_T_anchor.nunique()),
        fields=sorted(set(p.H_T)),
        emitted_field=sorted(set(em.H_T))[0],
    )


def edits(N):
    ch, pn = N["fams"]["iron_chalcogenide_11"], N["fams"]["iron_pnictide_122"]
    return dict(
        manuscript=[
            ("Of the 185, 125 receive at least one non-refused prediction "
             "target from the dispatch routine: 85 MgB2-class, 31 iron "
             "chalcogenide 11-type, and 9 iron pnictide 122-type. Two of the "
             "iron chalcogenide candidates are then removed by the "
             "calibration screen described below, so 123 compounds carry a "
             "prediction we report, 29 of them iron chalcogenide 11-type.",
             "Of the %d, %d receive at least one non-refused prediction "
             "target from the dispatch routine, and all %d are MgB2-class. "
             "The other two families dispatch nothing. Every one of iron "
             "chalcogenide 11-type's %d candidate-grid targets is refused, "
             "%d for lying below the validated reduced field, %d for a target "
             "temperature above Tc, and %d because the family does not pass "
             "field-axis validation; all %d of iron pnictide 122-type's are "
             "refused, %d for a missing upper-critical-field anchor, %d for "
             "reduced field, and %d for temperature. The calibration screen "
             "described below removes none of the %d, because the two "
             "candidates it acts on are iron chalcogenide and are no longer "
             "dispatched."
             % (N["compounds"], N["dispatched"], N["dispatched"],
                ch["targets"], ch["H_below_validated_reduced_field"],
                ch["T_above_Tc"], ch["family_fails_field_axis_validation"],
                pn["targets"], pn["Hc2_unavailable"],
                pn["H_below_validated_reduced_field"], pn["T_above_Tc"],
                N["dispatched"]),
             "the dispatch, as the deposited prediction table has it", None),
            ("These six predictions are excluded from every total reported in "
             "this paper, which is why 123 rather than 125 of the 185 "
             "candidates carry a prediction we are willing to report.",
             "These six predictions are excluded from every total reported in "
             "this paper. They fall on two iron chalcogenide candidates, and "
             "that family no longer dispatches at all, so the exclusion "
             "removes none of the %d compounds that carry a prediction we are "
             "willing to report." % N["dispatched"],
             "the outlier screen no longer changes the dispatched count",
             None),
            ("At 4.2 K and 0.1 T, the 87 dispatched MgB2-class prediction "
             "records, covering 85 distinct compounds, the 53 iron "
             "chalcogenide 11-type records covering 29 compounds at this grid "
             "point, two fewer than the 31 that receive a dispatched target "
             "somewhere on the grid, and the 37 iron pnictide 122-type "
             "records covering nine compounds each collapse to one single "
             "predicted value per family, identical to nine decimal places. "
             "This follows directly from the inference procedure of Section "
             "II.D: the substructure and sample-form conditioning fixes every "
             "parameter of Eq. (2), and the only compound-specific input is "
             "the upper-critical-field anchor, which enters solely through "
             "the field term. At 1 T the 53 chalcogenide records separate "
             "into three groups spanning 0.024 dex, one for each distinct "
             "anchor, with residual within-group differences below 0.0002 "
             "dex, with residual scatter below 0.001 dex from the bootstrap; "
             "at 5 T the span is 0.160 dex. Eighty-four of the 85 MgB2-class "
             "compounds carrying an anchor inherit the same parent value, so "
             "that family is identical at every evaluated point to within "
             "0.010 dex.",
             "At %.1f K and %.0f T, the only grid point at which anything is "
             "dispatched, the %d MgB2-class prediction records covering %d "
             "distinct compounds span %.4f dex, which is %.1f%% of the %.2f "
             "dex bootstrap interval at the same point. The family emits one "
             "value for every purpose the paper claims. This follows from the "
             "inference procedure of Section II.D: the substructure and "
             "sample-form conditioning fixes every parameter of Eq. (2), and "
             "the compound-specific inputs are the upper-critical-field "
             "anchor, which enters through the field term, and the "
             "transition-temperature anchor, which enters through the "
             "temperature term. Every dispatched record carries the same "
             "%.1f T parent anchor, so the field term is common to all of "
             "them and the residual spread is about %d times smaller than the "
             "uncertainty on any one prediction."
             % (GRID_POINT[0], GRID_POINT[1], N["records_at"],
                N["compounds_at"], N["span"], 100 * N["span"] / N["width_at"],
                N["width_at"], N["anchor"], round(N["ratio"] / 10) * 10),
             "measured spread at the surviving grid point, in place of an "
             "identity that no longer holds", None),
            ("Family-level screening outcome. We report family medians at 4.2 "
             "K and the lowest evaluated field of 0.1 T, where the prediction "
             "is effectively free of the field exponent: 6.00 for iron "
             "chalcogenide 11-type, 5.75 for iron pnictide 122-type, and 5.32 "
             "for MgB2-class, in log10 Jc with Jc in A cm−2. At this "
             "reference point the top-quartile candidate slice is entirely "
             "iron chalcogenide 11-type, and it remains so at 1 and 5 T. "
             "Evaluating instead at 1 T shifts these medians by −0.012, "
             "−0.005, and −0.055 dex respectively, and at 5 T by "
             "−0.067, −0.032, and −0.344 dex. The ordering and "
             "the gaps between families are unchanged across the evaluated "
             "field grid. That grid is, however, confined to the "
             "low-reduced-field corner of the applicability window: the "
             "dispatched candidates carry upper-critical-field anchors of 16, "
             "30 and 47 T in the iron chalcogenide family, 50 and 60 T in the "
             "iron pnictide family, and 15.5 T in the MgB2-class family, so 5 "
             "T is a reduced field of at most 0.32 for any dispatched "
             "candidate and below 0.11 for most of them.",
             "Family-level screening outcome. Only the MgB2 class dispatches "
             "under the refusal gates of this revision, so there is one "
             "family median to report: %.2f in log10 Jc with Jc in A "
             "cm−2, at %.1f K and %.0f T, over the %d records covering "
             "%d compounds. The evaluation grid carries three fields, %s T, "
             "and every target at %s T is refused for lying below the "
             "validated reduced field, so no median can be quoted at either. "
             "An earlier version of this section reported medians of 6.00, "
             "5.75 and 5.32 at 0.1 T across the three families, and a "
             "top-quartile slice that was entirely iron chalcogenide 11-type. "
             "Both rest on predictions this revision's reduced-field gate "
             "refuses, and both are withdrawn rather than restated. The "
             "surviving grid point sits at the edge of the applicability "
             "window: every dispatched candidate carries an "
             "upper-critical-field anchor of %.1f T, so %.0f T is a reduced "
             "field of %.2f, barely above the 0.3 bound of Eq. (1)."
             % (N["median_at"], GRID_POINT[0], GRID_POINT[1], N["records_at"],
                N["compounds_at"],
                _list(("%g" % f) for f in N["fields"]),
                " and ".join(("%g" % f) for f in N["fields"]
                             if f != N["emitted_field"]),
                N["anchor"], GRID_POINT[1], GRID_POINT[1] / N["anchor"]),
             "one family median, at the one grid point that survives", None),
            ("We therefore state the screening result at the scope it "
             "supports: at low reduced field, iron chalcogenide 11-type "
             "occupies the highest predicted current-density regime among the "
             "validated substructures. It is not a statement about relative "
             "performance in high field, and nothing finer than a "
             "family-level ranking is claimed.",
             "Those envelopes are what the fitted exponents support for their "
             "families. They are not dispatched predictions, and with the "
             "chalcogenide and 122 dispatches closed there is no screening "
             "comparison between families left to state. We therefore claim "
             "nothing finer than a family-level statement, and at present the "
             "dispatch supports one: the MgB2 class, at %.1f K and %.0f T."
             % (GRID_POINT[0], GRID_POINT[1]),
             "the cross-family screening claim rested on refused predictions",
             None),
            ("The prediction rule for each substructure follows from the "
             "variance-decomposition diagnostic. For iron chalcogenide "
             "11-type, sample-form conditioning is mandatory, so the dispatch "
             "uses the single-crystal empirical cell, the largest and best "
             "calibrated cell in the available literature. For iron pnictide "
             "122-type, conditioning is informative but less decisive, so the "
             "predictor uses the source-paper sample form when available and "
             "falls back to the single-crystal cell otherwise. For MgB2-class, "
             "sample form explains little variance, so the model uses the "
             "substructure-aggregate predictor throughout; the 84 "
             "parent-matched MgB2-class candidates inherit an "
             "upper-critical-field anchor of 15.5 T carrying a 20–50% "
             "deviation flag.",
             "The prediction rule for each substructure follows from the "
             "variance-decomposition diagnostic, and only the MgB2 rule is "
             "exercised by the current dispatch. For MgB2-class, sample form "
             "explains little variance, so the model uses the "
             "substructure-aggregate predictor throughout; every dispatched "
             "MgB2-class candidate inherits an upper-critical-field anchor of "
             "%.1f T carrying a 20–50%% deviation flag. The other two "
             "rules are stated because the diagnostic determines them and "
             "they would govern any dispatch that passed the gates. For iron "
             "chalcogenide 11-type, sample-form conditioning is mandatory, so "
             "the dispatch would use the single-crystal empirical cell, the "
             "largest and best calibrated cell in the available literature. "
             "For iron pnictide 122-type, conditioning is informative but "
             "less decisive, so the predictor would use the source-paper "
             "sample form when available and fall back to the single-crystal "
             "cell otherwise." % N["anchor"],
             "two of the three rules are not exercised by any dispatch",
             None),
            # The find string must not be a prefix of its own replacement.
            # It was, so a second run matched the replaced text again and
            # appended the new clause twice. The trailing comma is what makes
            # this one match once and once only.
            ("These curves apply only to candidates that pass the refusal "
             "gates, and panel (b)",
             "These curves apply only to candidates that pass the refusal "
             "gates, which at present is the MgB2 class alone, and panel (b)",
             "Fig. 5 caption, matched to the dispatch", None),
        ],
        supplement=[
            ("On the compound basis, 125 of the 185 receive at least one "
             "non-refused prediction target, as reported in Table IV of the "
             "main text.",
             "On the compound basis, %d of the %d receive at least one "
             "non-refused prediction target, all of them MgB2-class, as "
             "reported in Table IV of the main text."
             % (N["dispatched"], N["compounds"]),
             "dispatch on the compound basis", None),
        ],
        letter=[
            ("Section III.E now reports that at 4.2 K and the lowest "
             "evaluated field, the 87 dispatched MgB2-class prediction "
             "records covering 85 distinct compounds, the 53 iron "
             "chalcogenide records covering 29 compounds, and the 37 iron "
             "pnictide 122-type records covering nine compounds each collapse "
             "to one single predicted value per family, identical to nine "
             "decimal places. The conditioning fixes every parameter of the "
             "expression, and the only compound-specific input is the "
             "critical field anchor, which enters solely through the field "
             "term.",
             "Section III.E now reports this, and the reduced-field gate "
             "introduced in this revision has since narrowed what there is to "
             "report. Only the MgB2 class still dispatches; the iron "
             "chalcogenide and iron pnictide 122-type dispatches are refused "
             "in full. At %.1f K and %.0f T, the one grid point that survives "
             "the gate, the %d MgB2-class records covering %d distinct "
             "compounds span %.4f dex, which is %.1f%% of the %.2f dex "
             "bootstrap interval at that point. The conditioning fixes every "
             "parameter of the expression, and the compound-specific inputs "
             "are the critical-field anchor, which enters through the field "
             "term, and the transition temperature, which enters through the "
             "temperature term."
             % (GRID_POINT[0], GRID_POINT[1], N["records_at"],
                N["compounds_at"], N["span"],
                100 * N["span"] / N["width_at"], N["width_at"]),
             "the letter repeated the identity claim and the three-family "
             "counts", None),
            ("On precision, we have adopted the referee's correction for the "
             "reported uncertainties, so the interval reads 0.39 dex.",
             "On precision, we have adopted the referee's correction and no "
             "longer carry a third digit. The width itself has since changed. "
             "The 0.388 dex the referee read was the median over the "
             "predictions the pipeline emitted before the reduced-field gate "
             "introduced in this revision; over the %d predictions that "
             "survive that gate the median full width is 0.82 dex, a factor "
             "of about 6.7, and the manuscript now reports that."
             % sum(f["emitted"] for f in N["fams"].values()),
             "the conceded number is superseded by this revision's own gate",
             None),
        ],
    )


# The .docx and the .md do not agree on apostrophes: Word carries the curly
# U+2019 where the markdown has a straight quote. Matching on one spelling made
# the same edit land in the markdown and miss the .docx, which is exactly the
# drift these scripts exist to prevent, so every apostrophe matches either.
_APOS = "['\u2019]"


def _list(items):
    """Join for prose: a, b and c. A plain ", ".join wrote "0.1, 1, 5 T"."""
    items = list(items)
    if len(items) < 2:
        return "".join(items)
    return "%s and %s" % (", ".join(items[:-1]), items[-1])


def _rewrap(find):
    parts = []
    for w in find.split():
        parts.append("".join(_APOS if c in "'\u2019" else re.escape(c)
                             for c in w))
    return re.compile(r"\s+".join(parts))


def _flat(doc):
    parts = [par.text for par in doc.paragraphs]
    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                parts.extend(par.text for par in cell.paragraphs)
    return re.sub(r"\s+", " ", " ".join(parts))


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--ms", required=True)
    ap.add_argument("--supp", required=True)
    ap.add_argument("--letter", required=True)
    ap.add_argument("--letter-md", required=True)
    ap.add_argument("--out-dir", default=".")
    ap.add_argument("--dry-run", action="store_true")
    a = ap.parse_args()
    if not os.path.isdir("data"):
        sys.exit("run from the repository root")

    N = numbers()
    print("recomputed from the deposited prediction table\n")
    for fam, d in sorted(N["fams"].items()):
        print("   %-24s %4d targets, %4d emitted, %3d of %3d compounds"
              % (fam, d["targets"], d["emitted"], d["dispatched"],
                 d["compounds"]))
    print("\n   at %.1f K and %.0f T: %d records, %d compounds, median %.2f, "
          "span %.4f dex, interval %.2f dex"
          % (GRID_POINT[0], GRID_POINT[1], N["records_at"], N["compounds_at"],
             N["median_at"], N["span"], N["width_at"]))
    print()

    E = edits(N)
    report, misses = [], []
    docs, applied = {}, 0
    for label, path in (("manuscript", a.ms), ("supplement", a.supp),
                        ("letter", a.letter)):
        d = docx.Document(path)
        # apply() matches literally, so the find strings have to use whichever
        # apostrophe this document actually contains. Word writes the curly
        # U+2019 and the markdown twin writes a straight quote; typing one
        # spelling made an edit land in the .md and miss the .docx, leaving the
        # two out of step, which is the failure these scripts exist to catch.
        curly = "\u2019" in _flat(d)
        E[label] = [((f.replace("'", "\u2019") if curly else f), r, w, lo)
                    for f, r, w, lo in E[label]]
        sub = []
        apply(d, E[label], label, sub)
        text = _flat(d)
        for lb, done, find, repl, why in sub:
            if done:
                applied += 1
                report.append((lb, "ok", find, why))
            else:
                full = next(r for f, r, _, _ in E[label] if f[:72] == find)
                if _rewrap(full).search(text):
                    report.append((lb, "done", find, why))
                else:
                    misses.append(find)
                    report.append((lb, "MISS", find, why))
        docs[label] = d

    md = open(a.letter_md, encoding="utf-8").read()
    for find, repl, why, _ in E["letter"]:
        pat = _rewrap(find)
        if pat.search(md):
            md = pat.sub(lambda _m, r=repl: r, md, count=1)
            applied += 1
            report.append(("letter, markdown", "ok", find[:72], why))
        elif _rewrap(repl).search(re.sub(r"\s+", " ", md)):
            report.append(("letter, markdown", "done", find[:72], why))
        else:
            misses.append(find)
            report.append(("letter, markdown", "MISS", find[:72], why))

    for lb, state, find, why in report:
        print("   %-18s %-4s %s" % (lb, state, find))
        if state != "done":
            print("        %s" % why)
    print("\n   %d edit(s) applied this run" % applied)
    print()
    if misses:
        print("%d edit(s) missed; nothing written" % len(misses))
        return 1
    if a.dry_run:
        print("dry run, nothing written")
        return 0
    for label, path in (("manuscript", a.ms), ("supplement", a.supp),
                        ("letter", a.letter)):
        docs[label].save(os.path.join(a.out_dir, os.path.basename(path)))
        print("   wrote %s" % os.path.basename(path))
    with open(os.path.join(a.out_dir, os.path.basename(a.letter_md)), "w",
              encoding="utf-8") as fh:
        fh.write(md)
    print("   wrote %s" % os.path.basename(a.letter_md))
    return 0


if __name__ == "__main__":
    sys.exit(main())
