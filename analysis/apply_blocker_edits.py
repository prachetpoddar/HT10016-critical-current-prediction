#!/usr/bin/env python3
"""Write the seven recomputed statistics into the three documents.

Source is out/*_repaired.docx, the versions carrying the corrected census and
the disclosure. Output goes to out_blockers/.

Four of the seven passages are rewritten rather than renumbered, because the
number alone would say the wrong thing:

  the temperature leave-one-compound-out errors fall by up to a factor of six
  against a threshold that is absolute, while the repairs compress the
  exponents the threshold is applied to;

  the field-scale exposure rises rather than staying flat, and an unqualified
  swap of 0.80 for 0.88 beside a smaller cohort reads as noise;

  the two Stage 2 errors are computed on cohorts that are 73 percent iron-based
  and 89 percent MgB2, so setting them side by side without saying so reports a
  material-class difference as a degradation;

  "converges" is not what the Form 3 count measures. The fitter never reads a
  convergence flag; the gate is at least 20 points and three distinct
  temperatures.

The values come from analysis/recompute_blocker_statistics.py and
analysis/rerun_closed_form_without_withdrawn.py, and the reasoning is recorded
in audit/blocker_statistics_20260905.md.

One number in the Table S1 paragraph is corrected on the way past. It says 14
of the fully fittable compounds are represented by papers whose per-paper fits
also pass physicality. That does not reproduce on any cohort: the published 23
intersect the provenance table in 13 compounds and the still-contributing rows
in 10. It is set to 10, the count on the repaired cohort, and the discrepancy
is recorded rather than carried.

    python analysis/apply_blocker_edits.py --dry-run
    python analysis/apply_blocker_edits.py --out-dir out_blockers

Run from the repository root.
"""
import argparse
import os

import docx

SRC = "out"
MS = "HT10016_revised_repaired.docx"
SUPP = "SUPPLEMENTAL_MATERIAL_revised_repaired.docx"
RESP = "RESPONSE_TO_REFEREES_repaired.docx"

EXPOSURE_OLD = ("The median ratio of measured maximum to assigned scale is "
                "0.80, and for 15 of 94 curves it exceeds 0.9.")
EXPOSURE_NEW = (
    "The median ratio of measured maximum to assigned scale is 0.88, and for "
    "11 of the 52 curves admitted by the stated fitting protocol it exceeds "
    "0.9. On the cohort as published, before the anchor repairs, the same rule "
    "gave 0.80 and 15 of 94; the exposure rises rather than falls, from 16% of "
    "curves to 21%.")

MS_EDITS = [
    # blocker 1
    ("Using the 260 temperature-axis fits",
     "Using the 257 temperature-axis fits that survive the "
     "transition-temperature repairs",
     "the temperature cohort"),
    ("Iron chalcogenide 11-type reaches 0.588 across five compounds and 89 "
     "fits, with every scorable bootstrap resample below the screening-grade "
     "threshold of 1 in the exponent.",
     "Iron chalcogenide 11-type reaches 0.546 across three compounds and 87 "
     "fits, with every scorable bootstrap resample below the screening-grade "
     "threshold of 1 in the exponent.",
     "chalcogenide row"),
    ("Iron pnictide 122-type reaches 1.314 across two compounds, with 30% of "
     "resamples below threshold. Iron pnictide 1111-type reaches 3.120 across "
     "three compounds, with 21% below threshold. Only iron chalcogenide "
     "11-type clears the screening-grade threshold on this axis with "
     "confidence.",
     "Iron pnictide 122-type reaches 0.580 across two compounds, with 98% of "
     "resamples below threshold. Iron pnictide 1111-type reaches 0.513 across "
     "three compounds, with 100% below threshold. On the cohort as published "
     "these three errors were 0.588, 1.314 and 3.120, with 100%, 30% and 21% "
     "of resamples below threshold, and only the chalcogenide family cleared "
     "the threshold with confidence. The improvement has to be read with care. "
     "The threshold is an absolute bar of 1 in the exponent, and the repairs "
     "compress the exponents themselves: the mean absolute deviation of the "
     "1111-type exponents from their family median falls from 1.456 to 0.438, "
     "and predicting every fit by that median with nothing held out already "
     "gives an error equal to that deviation, so a family can cross the bar "
     "without the predictor having improved. Two things indicate that is not "
     "the whole of it. The 122-type family sits below 1 in both cohorts and "
     "still moves from 30% to 98%, and the error divided by the family's own "
     "deviation falls in both moving families, from 1.37 to 1.24 and from "
     "2.14 to 1.17. A third point runs the other way: the chalcogenide family "
     "loses two of its five compounds, both single fits, so its unchanged "
     "figure compares a five-fold estimator with a three-fold one. We report "
     "the errors on the repaired cohort with the compression stated beside "
     "them rather than claiming an improvement in the ratio of the two.",
     "122 and 1111 rows plus the compression caveat"),
    # blocker 2
    ("Ba(FeAs)2 with 85 fits and K(FeAs)2 with 21",
     "Ba(FeAs)2 with 84 fits and K(FeAs)2 with 21",
     "122 family composition"),
    ("the 260-fit temperature-exponent cohort contains no MgB2 fits",
     "the 257-fit temperature-exponent cohort contains no MgB2 fits, as the "
     "260-fit cohort before the repairs did not",
     "the MgB2 statement"),
    # the fully fittable cohort, wherever a statistic computed on it is
    # quoted. These are not renumbered to 20: the leave-one-out ranking, the
    # 81-cell scaling test, the 400-draw perturbation and the manual
    # verification were all run on the 23, and putting 20 next to them would
    # be the error this exercise exists to remove.
    ("Each of the 23 fully fittable compounds is held out in turn",
     "Each of the 23 fully fittable compounds of the pre-withdrawal cohort is "
     "held out in turn",
     "Fig. 4 leave-one-out"),
    ("The test bins the 23 fully fittable compounds on an 81-cell grid",
     "The test bins the 23 fully fittable compounds of the pre-withdrawal "
     "cohort on an 81-cell grid",
     "the universal-scaling test"),
    ("manual verification by the authors against the source figures for all "
     "23 fully fittable compounds",
     "manual verification by the authors against the source figures for all "
     "23 fully fittable compounds of the pre-withdrawal cohort",
     "the manual verification"),
    ("of which 23 compounds satisfy the stricter two-axis criterion",
     "of which 20 satisfy the stricter two-axis criterion, 23 before the "
     "eleven withdrawals",
     "Fig. 1 caption"),
    # blocker 3
    ("the median ratio of measured maximum to assigned scale is 0.80, and for "
     "15 of 94 curves it exceeds 0.9",
     "the median ratio of measured maximum to assigned scale is 0.88, and for "
     "11 of the 52 curves admitted by the stated fitting protocol it exceeds "
     "0.9, against 0.80 and 15 of 94 on the cohort as published, so the "
     "exposure rises rather than falls",
     "the exposure ratio, main text"),
]

MS_TABLE_EDITS = [
    (0, 5, 1, "23", "20", "Table I, fully fittable compounds"),
    (0, 5, 0, "Fully fittable compounds (both axes)",
     "Compounds with a computable aggregate fit (both axes)",
     "Table I, the row label"),
    (0, 5, 2, "Functional-form selection; universal-scaling test",
     "Functional-form selection; universal-scaling test. Both were computed "
     "on the pre-withdrawal set of 23 and are reported as such where they "
     "appear",
     "Table I, what row 5 supports"),
    # The leave-one-compound-out summary row. Its temperature clause is the
    # same statistic as Sec. III.C and moves with it. Its field clause is not
    # renumbered: the repaired cohort reverses which families clear the
    # threshold, and that reversal propagates into Table III's per-family
    # verdicts and into the dispatch claims that rest on them. Renumbering it
    # here without carrying that through would leave the paper asserting two
    # different dispatch scopes, so the cell says plainly that the field
    # clause is on the published cohort and that the repaired one disagrees.
    (2, 4, 2,
     "Temperature axis: iron chalcogenide 11-type 0.588 clears the "
     "screening-grade threshold of 1 in the exponent; 122-type 1.314 and "
     "1111-type 3.120 do not; MgB2-class not assessable. Field axis: "
     "MgB2-class 0.753 and 122-type 0.973 clear; 11-type 1.093 and 1111-type "
     "2.571 do not.",
     "Temperature axis, on the repaired 257-fit cohort: iron chalcogenide "
     "11-type 0.546, 122-type 0.580 and 1111-type 0.513 all fall below the "
     "screening-grade threshold of 1 in the exponent, against 0.588, 1.314 "
     "and 3.120 as published; the threshold is absolute and the repairs "
     "compress the exponents, so Sec. III.C states what part of that movement "
     "is scale rather than skill. MgB2-class not assessable. Field axis, on "
     "the cohort as published: MgB2-class 0.753 and 122-type 0.973 clear; "
     "11-type 1.093 and 1111-type 2.571 do not. On the repaired 52-fit cohort "
     "those become 1.230, 1.957, 0.707 and 3.327, which reverses which "
     "families clear. That reversal is not yet carried through the "
     "family-level dispatch scope stated elsewhere in this paper, and the "
     "field-axis verdicts in this row remain those of the published cohort "
     "until it is.",
     "Table II, the leave-one-out summary row"),
    (2, 1, 1, "binning across the 23 fully fittable compounds",
     "binning across the 23 fully fittable compounds of the pre-withdrawal "
     "cohort", "Table II, the universal-scaling row"),
]

SUPP_EDITS = [
    ("Across the 23 fully fittable compounds, the one cohort in which all "
     "three parameters are determined",
     "Across the 23 fully fittable compounds of the pre-withdrawal cohort, "
     "the one cohort in which all three parameters are determined",
     "the correlation argument"),
    # blocker 3
    (EXPOSURE_OLD, EXPOSURE_NEW, "the exposure ratio, supplement"),
    # blocker 4
    ("The Stage 2 predictor, conditioned on substructure and sample form, "
     "gives a mean absolute error of 1.257 in βH over 82 scored fits from "
     "16 papers, with a 95% bootstrap confidence interval of [1.006, 1.612]. "
     "A pooled median with no conditioning gives 1.158 over 94 fits from the "
     "same 16 papers, with [0.900, 1.497]. The bootstrap uses 5000 iterations "
     "with seed 20260901 and the percentile method throughout. The two are "
     "not scored on the same cohort, since the conditioned predictor cannot "
     "score a fit whose sample-form cell is unpopulated, so this is not a "
     "like-for-like comparison; but the conditioned predictor no longer has "
     "the lower error and we state that rather than reporting only the figure "
     "that flatters it.",
     "The Stage 2 predictor, conditioned on substructure and sample form, "
     "gives a mean absolute error of 1.572 in βH over the 19 fits it can "
     "score out of the 52 admitted, drawn from 6 papers, with a 95% bootstrap "
     "confidence interval of [0.800, 2.253]. A pooled median with no "
     "conditioning gives 1.361 over all 52 fits from 12 papers, with [0.741, "
     "2.130]. On the cohort as published the same two predictors gave 1.257 "
     "over 82 fits and 1.158 over 94, from 16 papers. The bootstrap uses 5000 "
     "iterations with seed 20260901 and the percentile method throughout, and "
     "resamples source papers rather than individual residuals; resampling "
     "residuals independently treats several readings of one figure as "
     "independent observations and returns [1.201, 1.958] and [1.039, 1.789], "
     "which are too narrow to support any statement about the threshold. The "
     "two Stage 2 figures are not the same statistic on a smaller cohort. "
     "Stage 2 cannot score a fit whose combination of substructure and sample "
     "form is represented by no other paper, which on the published cohort "
     "leaves 82 of 94 fits, 73% of them iron-based, and on the repaired "
     "cohort leaves 19 of 52, 89% of them MgB2: every iron chalcogenide and "
     "1111-type cell is empty and the whole iron contribution is two "
     "thin-film 122-type fits. The move from 1.257 to 1.572 therefore sets "
     "the error on one material class beside the error on another. What the "
     "repairs changed is Stage 2's reach, from 87% of the cohort to 37%. The "
     "pooled row, which scores every fit, is the one comparable across the "
     "two cohorts. The conditioned predictor no longer has the lower error on "
     "either cohort, and we state that rather than reporting only the figure "
     "that flatters it.",
     "the per-paper validation"),
    # blocker 5
    ("the paper’s compound is among the 23 compounds whose per-compound "
     "aggregate Form 3 fit converges; 14 of these compounds are represented "
     "by papers whose per-paper fits also pass physicality and therefore "
     "appear here",
     "the paper’s compound is among the 20 compounds, of the 24 that "
     "retain any data after the eleven withdrawals, whose per-compound "
     "aggregate Form 3 fit is computable; 10 of these are represented by "
     "papers whose per-paper fits also pass physicality and therefore appear "
     "here",
     "the fully fittable definition"),
    ("or Cohort A and B with a non-fittable aggregate compound (both axes "
     "contributed, but the per-compound aggregate fit does not converge at "
     "the current cohort scope)",
     "or Cohort A and B with a non-fittable aggregate compound (both axes "
     "contributed, but the per-compound aggregate fit is not computable at "
     "the current cohort scope). Computable here means that at least 20 "
     "measurements and three distinct temperatures survive the Form 3 "
     "inclusion filter. An earlier version of this section called it "
     "convergence, which the underlying fitter does not test: it returns the "
     "least-squares parameters without reading a convergence flag, and "
     "several of the compounds it admits are poorly constrained. On the "
     "cohort as published the count was 23 of 27",
     "the non-fittable definition and the convergence wording"),
]

RESP_EDITS = [
    ("The median ratio of measured maximum to assigned scale is 0.80, and 15 "
     "of 94 curves sit above 0.9.",
     "The median ratio of measured maximum to assigned scale is 0.88, and 11 "
     "of the 52 curves admitted by the stated fitting protocol sit above 0.9. "
     "On the cohort as published the same rule gave 0.80 and 15 of 94, so the "
     "exposure rises rather than falls once the anchors are repaired.",
     "the exposure ratio, response"),
]


def replace_in_paragraph(p, find, repl):
    """Replace across runs, keeping the first run's formatting."""
    text = "".join(r.text for r in p.runs)
    if find not in text:
        return False
    start = text.index(find)
    end = start + len(find)
    pos, first, spans = 0, None, []
    for r in p.runs:
        a, b = pos, pos + len(r.text)
        if b > start and a < end:
            spans.append((r, max(start, a) - a, min(end, b) - a))
            if first is None:
                first = r
        pos = b
    if first is None:
        return False
    for r, i, j in spans:
        r.text = r.text[:i] + (repl if r is first else "") + r.text[j:]
    return True


def apply_doc(path, edits, table_edits=()):
    d = docx.Document(path)
    missed = []
    for find, repl, why in edits:
        if not any(replace_in_paragraph(p, find, repl) for p in d.paragraphs):
            missed.append((why, find[:70]))
    for ti, row, col, find, repl, why in table_edits:
        cell = d.tables[ti].rows[row].cells[col]
        if not any(replace_in_paragraph(p, find, repl)
                   for p in cell.paragraphs):
            missed.append((why, find[:70]))
    return d, missed


STALE = ["15 of 94", "over 94 fits", "1.158 over", "260 temperature-axis",
         "260-fit temperature-exponent", "85 fits and K(FeAs)2",
         "23 compounds whose per-compound"]

# A superseded figure may stay in the text as long as the sentence says it is
# superseded. Several of the rewrites above quote the published value on
# purpose, so a bare string search would fire on the fix itself. The test is
# whether the paragraph carrying the number also says which cohort it belongs
# to.
QUALIFIERS = ["as published", "before the anchor repairs", "before the repairs",
              "An earlier version", "internally inconsistent",
              "as the 260-fit cohort"]


def check_residue(d, name):
    """Numbers that belonged to a cohort that no longer exists, unqualified.

    Run after the edits. Anything still present without a qualifier is either
    an edit that landed somewhere unexpected or a passage nobody has looked at.
    """
    hits = []
    blocks = [p.text for p in d.paragraphs]
    for t in d.tables:
        for r in t.rows:
            blocks += [c.text for c in r.cells]
    for text in blocks:
        if any(q in text for q in QUALIFIERS):
            continue
        for s in STALE:
            if s in text:
                hits.append(s)
    if hits:
        print("   residue in %s: %s" % (name, ", ".join(sorted(set(hits)))))
    return hits


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--out-dir")
    ap.add_argument("--dry-run", action="store_true")
    a = ap.parse_args()

    docs = [(MS, MS_EDITS, MS_TABLE_EDITS), (SUPP, SUPP_EDITS, ()),
            (RESP, RESP_EDITS, ())]
    results, all_missed, residue = [], [], []
    for name, edits, tedits in docs:
        d, missed = apply_doc(os.path.join(SRC, name), edits, tedits)
        results.append((name, d))
        all_missed += [(name,) + m for m in missed]
        print("%s: %d edits, %d not found"
              % (name, len(edits) + len(tedits), len(missed)))
        residue += check_residue(d, name)
    if all_missed:
        print("\nEDITS THAT FOUND NO TARGET. Nothing written.")
        for m in all_missed:
            print("  ", m)
        return 1
    if residue:
        print("\nSTALE NUMBERS STILL PRESENT. Nothing written.")
        return 1
    if a.dry_run or not a.out_dir:
        print("\ndry run: nothing written")
        return 0
    os.makedirs(a.out_dir, exist_ok=True)
    for name, d in results:
        out = os.path.join(a.out_dir, name.replace("_repaired", "_final2"))
        d.save(out)
        print("written: %s" % out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
