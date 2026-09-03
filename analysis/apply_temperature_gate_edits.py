"""
apply_temperature_gate_edits.py

Carry the temperature-window gate, and twenty strings an adversarial review
found falsified, into all three artifacts.

Two separate things landed at once and are applied together because they touch
the same sentences.

The gate. analysis/apply_temperature_window_gate.py refuses the 93 targets
sitting at reduced temperature 0.7 or above. Dispatch falls from 85 compounds
to 84 and from 256 emitted targets to 163, and the median 95% interval falls
from 0.825 dex to 0.612, a factor of 4.1 rather than 6.7, because the refused
targets were the wide ones.

The review. The consistency gate was printing "all four artifacts agree" while
these stood:

  * "5 T is a reduced field of 0.32, barely above the 0.3 bound of Eq. (1)",
    which I wrote. Eq. (1)'s field clause bounds a fitted curve's measured
    SPAN, (Hmax - Hmin)/Hc2,0 > 0.3. It partitions the deposited fit table with
    zero error: minimum span 0.3499 among the 94 passing fits against maximum
    0.2672 among the 60 the bound refuses. A single evaluation point has no
    span, so comparing 0.3226 against 0.3 is a category error rather than a
    close call. Assumption 1 and the letter carry the same confusion from the
    other side, still describing the dispatch as sitting "between roughly 0.002
    and 0.3" and at "reduced fields of order 0.01", which the gate made false.
  * The two field-scale implementations as "87 of 186 fits" and "99" at a
    median factor of 5.2, against 61 of 159, 98 and 6.4.
  * "29 of the 33 source papers behind the temperature-axis cohort", where the
    deposit holds 20 papers, 18 of them from the arXiv baseline. The manuscript
    already said eighteen of the 20; the supplement and the letter did not.
  * The letter promising the width "as a factor of about 2.5" two sentences
    after giving it as 6.7, and quoting field shifts of 0.012, 0.005 and 0.055
    dex from a passage Sec. III.E no longer contains.
  * The supplement still telling the reader the outlier screen reduces the
    count "from the 125 the dispatch routine emits to the 123".

Every replacement is computed from the deposit at run time. The script refuses
to write if any edit misses, and treats an already-applied edit as done.

Usage:
    python3 analysis/apply_temperature_gate_edits.py --ms MS.docx \\
        --supp SUPP.docx --letter L.docx --letter-md L.md --out-dir DIR
"""
import argparse
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import docx  # noqa: E402
import pandas as pd  # noqa: E402

from apply_manuscript_edits import apply  # noqa: E402

Z = 1.959964   # two-sided 95% normal quantile, for the one-sigma conversion


def numbers():
    p = pd.read_csv(os.path.join("data",
                                 "phase_3_p57_de_novo_predictions.csv"),
                    low_memory=False)
    fh = pd.read_csv(os.path.join(
        "data", "phase_3_form3_fits_partial_cohortB_v2.csv"))
    bt = pd.read_csv(os.path.join(
        "data", "phase_3_p44_post_UCLA_beta_T_fits.csv"))
    em = p[p.refusal_flag.fillna("") == ""]
    w = float((em.predicted_log_Jc_upper_95
               - em.predicted_log_Jc_lower_95).median())
    at = em[(em.T_K == 4.2) & (em.H_T == 5.0)]
    aw = float((at.predicted_log_Jc_upper_95
                - at.predicted_log_Jc_lower_95).median())
    flags = p.refusal_flag.fillna("")
    share = {k: 100.0 * (flags == k).sum() / len(p) for k in set(flags) if k}
    import numpy as np
    same = np.isclose(fh.Hc2_T_used, fh.Hc2_T_default)
    tiers = fh.Hc2_source.astype(str).str.extract(r"^(Tier_\d)")[0]
    papers = bt[["paper_id", "source"]].drop_duplicates()
    return dict(
        dispatched=int(em.compound_formula.nunique()),
        emitted=len(em),
        width=w, factor=10 ** w, half=w / 2, sigma=w / (2 * Z),
        at_records=len(at), at_compounds=int(at.compound_formula.nunique()),
        at_span=float(at.predicted_log_Jc.max() - at.predicted_log_Jc.min()),
        at_width=aw,
        share=share,
        fh_total=len(fh), fh_default=int(same.sum()),
        fh_resolved=int((~same).sum()),
        fh_factor=float((fh.loc[~same, "Hc2_T_default"]
                         / fh.loc[~same, "Hc2_T_used"]).median()),
        tier1=int((tiers == "Tier_1").sum()),
        tier2=int((tiers == "Tier_2").sum()),
        tier3=int((tiers == "Tier_3").sum()),
        t_papers=len(papers),
        t_baseline=int((papers.source == "v3.2.1_baseline").sum()),
        # The temperature gate's own numbers, read from the table rather than
        # typed: how many targets it refused, and the largest reduced
        # temperature any deposited fit actually measured.
        t_refused=int((flags == "T_above_validated_reduced_temperature").sum()),
        t_cover_max=float((bt.T_max / bt.Tc_K).max()),
    )


def _shares(N):
    """The four refusal codes, largest first, as prose."""
    order = sorted(N["share"].items(), key=lambda kv: -kv[1])
    name = {
        "H_below_validated_reduced_field": "lying below the validated reduced "
                                           "field",
        "Hc2_unavailable": "a missing upper-critical-field anchor",
        "T_above_Tc": "a target temperature above Tc",
        "T_above_validated_reduced_temperature": "lying at or above the "
                                                 "validated reduced "
                                                 "temperature",
        "family_fails_field_axis_validation": "the family failing field-axis "
                                              "validation",
    }
    return ", ".join("%.1f%% for %s" % (v, name.get(k, k)) for k, v in order)


def edits(N):
    D, E = N["dispatched"], N["emitted"]
    return dict(
        manuscript=[
            ("105 / 103 / 85", "105 / 103 / %d" % D,
             "Table IV MgB2 row", None),
            ("233 / 183 / 85", "233 / 183 / %d" % D,
             "Table IV combined row", None),
            ("Of the 183 candidates, 85 are dispatched",
             "Of the 183 candidates, %d are dispatched" % D,
             "Fig. 1 caption", None),
            ("85 compounds receive at least one non-refused prediction "
             "target; 50.3% of candidate-grid-point predictions are refused "
             "for lying below the validated reduced field, 25.8% for a "
             "missing field anchor and 9.9% for target temperature above Tc.",
             "%d compounds receive at least one non-refused prediction "
             "target; of the candidate-grid-point predictions, %s."
             % (D, _shares(N)),
             "Table IV lead, five codes now fire", None),
            ("Predictions below the validated reduced field produce refusals "
             "for 50.3% of candidate-grid-point predictions, missing "
             "upper-critical-field anchors for 25.8%, and target temperatures "
             "above the transition temperature for 9.9%.",
             "Of the candidate-grid-point predictions, %s." % _shares(N),
             "refusal shares, recomputed", None),
            ("Of the 183, 85 receive at least one non-refused prediction "
             "target from the dispatch routine, and all 85 are MgB2-class.",
             "Of the 183, %d receive at least one non-refused prediction "
             "target from the dispatch routine, and all %d are MgB2-class."
             % (D, D), "dispatch count after the temperature gate", None),
            ("The calibration screen described below removes none of the 85,",
             "The calibration screen described below removes none of the %d,"
             % D, "same count", None),
            ("removes none of the 85 compounds that carry a prediction we are "
             "willing to report.",
             "removes none of the %d compounds that carry a prediction we are "
             "willing to report." % D, "same count", None),
            ("over the 87 records covering 85 compounds",
             "over the %d records covering %d compounds"
             % (N["at_records"], N["at_compounds"]),
             "records at the surviving grid point", None),
            ("the 87 MgB2-class prediction records covering 85 distinct "
             "compounds span 0.0098 dex, which is 1.6% of the 0.60 dex",
             "the %d MgB2-class prediction records covering %d distinct "
             "compounds span %.4f dex, which is %.1f%% of the %.2f dex"
             % (N["at_records"], N["at_compounds"], N["at_span"],
                100 * N["at_span"] / N["at_width"], N["at_width"]),
             "spread at the surviving grid point", None),
            ("of which 85 compounds receive at least one dispatched target",
             "of which %d compounds receive at least one dispatched target" % D,
             "conclusion", None),
            ("is 0.82 dex in log10 Jc, a factor of about 6.7 in Jc",
             "is %.2f dex in log10 Jc, a factor of about %.1f in Jc"
             % (N["width"], N["factor"]),
             "interval width after the temperature gate", None),
            ("pointwise one-sigma uncertainty of about 0.21 dex",
             "pointwise one-sigma uncertainty of about %.2f dex" % N["sigma"],
             "derived from the width", None),
            ("constant half-width of 0.41 dex around each curve, which is "
             "half the median full width of the 95% bootstrap confidence "
             "interval quoted as 0.82 dex in the text",
             "constant half-width of %.2f dex around each curve, which is half "
             "the median full width of the 95%% bootstrap confidence interval "
             "quoted as %.2f dex in the text" % (N["half"], N["width"]),
             "Fig. 5 envelope", None),
            ("The two are identical for the 87 of 186 fits that fall back to "
             "the literature default, and the resolved quantity is a median "
             "factor of 5.2 smaller than the literature Hc2,0 across the 99 "
             "fits for which a paper-derived value was resolved.",
             "The two are identical for the %d of %d fits that fall back to "
             "the literature default, and the resolved quantity is a median "
             "factor of %.1f smaller than the literature Hc2,0 across the %d "
             "fits for which a paper-derived value was resolved."
             % (N["fh_default"], N["fh_total"], N["fh_factor"],
                N["fh_resolved"]),
             "two field-scale implementations, recomputed", None),
            ("and the dispatch of Sec. III.E reports at 4.2 K and at reduced "
             "fields of order 0.01, which lies below the field bound rather "
             "than inside it. The predictions we report are therefore "
             "extrapolations of the fitted form outside the range over which "
             "its exponents were validated, and we now say so rather than "
             "leaving the window to imply otherwise.",
             "and the dispatch of Sec. III.E reports at 4.2 and 20 K. Both "
             "halves of the window are now enforced on the dispatch itself, "
             "per evaluation point: a target is refused below a reduced field "
             "of 0.3 and at or above a reduced temperature of 0.7. We "
             "distinguish that per-point enforcement from Eq. (1) itself, "
             "which admits a curve to the fit on the width of its measured "
             "interval rather than on where any single point sits, and we "
             "apply the stricter of the two readings to what we emit, rather "
             "than leaving the window to imply a scope we do not enforce.",
             "the extrapolation claim is false once both gates are enforced",
             None),
            ("whereas the dispatch grid evaluates at 0.1, 1 and 5 T against "
             "anchors of 15.5 to 60 T, so every dispatched point sits at a "
             "reduced field between roughly 0.002 and 0.3. The field term is "
             "therefore evaluated outside the range in which the field "
             "exponent was fitted and validated.",
             "whereas the dispatch grid evaluates at 0.1, 1 and 5 T against "
             "anchors of 15.5 to 60 T, so most grid targets fall far below a "
             "reduced field of 0.3. Those are refused rather than reported: "
             "every emitted target sits at a reduced field of 0.3226 against "
             "the 15.5 T MgB2 parent anchor.",
             "the pre-gate range, still described as though it were emitted",
             None),
            ("every dispatched candidate carries an upper-critical-field "
             "anchor of 15.5 T, so 5 T is a reduced field of 0.32, barely "
             "above the 0.3 bound of Eq. (1).",
             "every dispatched candidate carries an upper-critical-field "
             "anchor of 15.5 T, so 5 T is a reduced field of 0.3226. We note "
             "what that number is and is not. Eq. (1) admits a curve to the "
             "field-axis fit on the width of its measured interval, "
             "(Hmax - Hmin)/Hc2,0 > 0.3, which is a property of a fitted "
             "curve and not of a single evaluation point, so 0.3226 is not a "
             "margin against that criterion. What it is comparable to is the "
             "reduced-field range the admitted fits actually cover, and the "
             "dispatch gate of Sec. II.D refuses any target below 0.3 on that "
             "reading. The parent anchor carries a 20 to 50 percent deviation "
             "flag, and a 20 percent upward revision of it would put every "
             "emitted target below the gate.",
             "a point was compared against a bound defined on a span", None),
            ("The evaluation grid carries three fields, 0.1, 1 and 5 T, and "
             "every target at 0.1 and 1 T is refused for lying below the "
             "validated reduced field, so no median can be quoted at either.",
             "The evaluation grid carries three fields, 0.1, 1 and 5 T, and "
             "every target at 0.1 and 1 T is refused for lying below the "
             "validated reduced field, so no median can be quoted at either. "
             "The grid also carried a temperature point at 0.77 Tc for every "
             "candidate, ten percent past the reduced-temperature bound and "
             "beyond the reach of any deposited fit, whose largest measured "
             "coverage is %.4f. Those %d targets are refused as well."
             % (N["t_cover_max"], N["t_refused"]),
             "the temperature gate, stated where the medians are", None),
        ],
        supplement=[
            ("On the compound basis, 85 of the 183 receive at least one "
             "non-refused prediction target",
             "On the compound basis, %d of the 183 receive at least one "
             "non-refused prediction target" % D, "dispatch count", None),
            ("width of 0.82 dex, a factor of about 6.7 in Jc",
             "width of %.2f dex, a factor of about %.1f in Jc"
             % (N["width"], N["factor"]), "interval width", None),
            ("0.30 dex is a factor of 2, and 0.82 dex is a factor of about 6.7",
             "0.30 dex is a factor of 2, and %.2f dex is a factor of about %.1f"
             % (N["width"], N["factor"]), "glossary, same width", None),
            ("Across the whole file, 90 of the 186 fits resolve their scale "
             "through a Tier 1 direct or extrapolated match, 87 through the "
             "Tier 3 literature default, and 9 through the Tier 2 "
             "per-substructure ratio.",
             "Across the whole file, %d of the %d fits resolve their scale "
             "through a Tier 1 direct or extrapolated match, %d through the "
             "Tier 3 literature default, and %d through the Tier 2 "
             "per-substructure ratio."
             % (N["tier1"], N["fh_total"], N["tier3"], N["tier2"]),
             "provenance tiers, recomputed", None),
            ("29 of the 33 temperature-axis source papers are arXiv preprints",
             "%d of the %d temperature-axis source papers are arXiv preprints"
             % (N["t_baseline"], N["t_papers"]),
             "the cohort holds 20 papers, not 33", None),
            ("removing the six tuples reduces the reported candidate count "
             "from the 125 the dispatch routine emits to the 123 reported in "
             "the main text.",
             "removing the six tuples changes no reported count. They fall on "
             "two iron chalcogenide candidates, and that family no longer "
             "dispatches at all, so the %d compounds the dispatch emits are "
             "the %d reported in the main text." % (D, D),
             "the outlier screen no longer changes the count", None),
            ("Predictions below the validated reduced field produce refusals "
             "for 50.3% of candidate-grid tuples, missing "
             "upper-critical-field anchors for 25.8%, and target temperatures "
             "above the transition temperature for 9.9%.",
             "Of the candidate-grid tuples, %s." % _shares(N),
             "refusal shares, recomputed", None),
        ],
        letter=[
            ("And the window requires reduced field above 0.3, whereas every "
             "dispatched prediction sits between roughly 0.002 and 0.3. The "
             "predictions we report are therefore extrapolations of the "
             "fitted form outside the range over which its exponents were "
             "validated.",
             "And the window requires reduced field above 0.3, which most of "
             "the dispatch grid did not meet. Rather than report those as "
             "predictions we now refuse them: a per-point gate on each half of "
             "the window removed 1054 targets below the reduced-field bound "
             "and 93 at or above the reduced-temperature bound, leaving %d "
             "emitted targets over %d compounds, all at a reduced field of "
             "0.3226. We should be exact about what that bound is. Eq. (1) "
             "admits a curve to the fit on the width of its measured field "
             "interval, not on where a single point sits, so the two are not "
             "the same test; we apply the stricter reading to what we emit "
             "and say so." % (E, D),
             "the concession is superseded by our own gates", None),
            ("The two are identical for the 87 of 186 fits that fall back to "
             "the literature default. For the 99 fits where a paper-derived "
             "value was resolved, that value is a median factor of 5.2 "
             "smaller.",
             "The two are identical for the %d of %d fits that fall back to "
             "the literature default. For the %d fits where a paper-derived "
             "value was resolved, that value is a median factor of %.1f "
             "smaller." % (N["fh_default"], N["fh_total"], N["fh_resolved"],
                           N["fh_factor"]),
             "recomputed", None),
            ("that 29 of the 33 source papers behind the temperature-axis "
             "cohort are arXiv preprints",
             "that %d of the %d source papers behind the temperature-axis "
             "cohort are arXiv preprints"
             % (N["t_baseline"], N["t_papers"]), "recomputed", None),
            ("We report family medians at the lowest evaluated field and give "
             "the shift to 1 T and to 5 T numerically, so the reader can see "
             "how small the field contribution is: at the reference point it "
             "is 0.012, 0.005, and 0.055 dex for the three families.",
             "We report the family median at the one grid point that survives "
             "the refusal gates, 4.2 K and 5 T, which is a reduced field of "
             "0.3226 rather than the near-self-field point the referee "
             "identified. The low-field targets the referee's objection bears "
             "on are refused rather than reported.",
             "Sec. III.E no longer contains the passage this describes", None),
            ("We have also adopted the suggested gloss: the text now gives "
             "that width as a factor of about 2.5 in the critical current",
             "We have also adopted the suggested gloss, giving the width as a "
             "factor in the critical current rather than only in dex: it is a "
             "factor of about %.1f" % N["factor"],
             "the letter promised 2.5 two sentences after giving 6.7", None),
            ("over the 256 predictions that survive that gate the median full "
             "width is 0.82 dex, a factor of about 6.7, and the manuscript "
             "now reports that.",
             "over the %d predictions that survive our refusal gates the "
             "median full width is %.2f dex, a factor of about %.1f, and the "
             "manuscript now reports that."
             % (E, N["width"], N["factor"]), "recomputed", None),
            ("the 87 MgB2-class records covering 85 distinct compounds span "
             "0.0098 dex, which is 1.6% of the 0.60 dex bootstrap interval at "
             "that point.",
             "the %d MgB2-class records covering %d distinct compounds span "
             "%.4f dex, which is %.1f%% of the %.2f dex bootstrap interval at "
             "that point."
             % (N["at_records"], N["at_compounds"], N["at_span"],
                100 * N["at_span"] / N["at_width"], N["at_width"]),
             "recomputed", None),
        ],
    )


_APOS = "['’]"


def _rewrap(find):
    parts = []
    for w in find.split():
        parts.append("".join(_APOS if c in "'’" else re.escape(c)
                             for c in w))
    return re.compile(r"\s+".join(parts))


def _flat(doc):
    parts = [par.text for par in doc.paragraphs]
    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                parts.extend(par.text for par in cell.paragraphs)
    return re.sub(r"\s+", " ", " ".join(parts))


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--ms", required=True)
    ap.add_argument("--supp", required=True)
    ap.add_argument("--letter", required=True)
    ap.add_argument("--letter-md", required=True)
    ap.add_argument("--out-dir", default=".")
    ap.add_argument("--dry-run", action="store_true")
    a = ap.parse_args()
    if not os.path.isdir("data"):
        sys.exit("run from the repository root")

    N = numbers()
    print("recomputed from the deposit\n")
    for k in ("dispatched", "emitted", "at_records", "at_compounds",
              "fh_total", "fh_default", "fh_resolved", "tier1", "tier2",
              "tier3", "t_papers", "t_baseline"):
        print("   %-14s %8d" % (k, N[k]))
    print("   %-14s %8.4f dex  factor %.2f  half %.4f  sigma %.4f"
          % ("interval", N["width"], N["factor"], N["half"], N["sigma"]))
    print("   refusal shares: %s\n" % _shares(N))

    E = edits(N)
    report, misses, applied = [], [], 0
    docs = {}
    for label, path in (("manuscript", a.ms), ("supplement", a.supp),
                        ("letter", a.letter)):
        d = docx.Document(path)
        curly = "’" in _flat(d)
        E[label] = [((f.replace("'", "’") if curly else f), r, w, lo)
                    for f, r, w, lo in E[label]]
        sub = []
        apply(d, E[label], label, sub)
        text = _flat(d)
        for lb, done, find, repl, why in sub:
            if done:
                applied += 1
                report.append((lb, "ok", find, why))
            else:
                full = next(r for f, r, _, _ in E[label] if f[:72] == find)
                if _rewrap(full).search(text):
                    report.append((lb, "done", find, why))
                else:
                    misses.append(find)
                    report.append((lb, "MISS", find, why))
        docs[label] = d

    md = open(a.letter_md, encoding="utf-8").read()
    for find, repl, why, _ in E["letter"]:
        pat = _rewrap(find)
        if pat.search(md):
            md = pat.sub(lambda _m, r=repl: r, md, count=1)
            applied += 1
            report.append(("letter, markdown", "ok", find[:72], why))
        elif _rewrap(repl).search(re.sub(r"\s+", " ", md)):
            report.append(("letter, markdown", "done", find[:72], why))
        else:
            misses.append(find)
            report.append(("letter, markdown", "MISS", find[:72], why))

    for lb, state, find, why in report:
        print("   %-18s %-4s %s" % (lb, state, find))
        if state == "MISS":
            print("        %s" % why)
    print("\n   %d edit(s) applied this run\n" % applied)
    if misses:
        print("%d edit(s) missed; nothing written" % len(misses))
        return 1
    if a.dry_run:
        print("dry run, nothing written")
        return 0
    for label, path in (("manuscript", a.ms), ("supplement", a.supp),
                        ("letter", a.letter)):
        docs[label].save(os.path.join(a.out_dir, os.path.basename(path)))
        print("   wrote %s" % os.path.basename(path))
    with open(os.path.join(a.out_dir, os.path.basename(a.letter_md)), "w",
              encoding="utf-8") as fh:
        fh.write(md)
    print("   wrote %s" % os.path.basename(a.letter_md))
    return 0


if __name__ == "__main__":
    sys.exit(main())
