"""
apply_reclassification_edits.py

Carry the corrected substructure classifier into the three artifacts.

analysis/fix_substructure_classifier.py moved four anchor rows out of
other_unclassified: two Cu-doped FeTe0.66Se0.34 crystals into iron chalcogenide
11-type, and two Ba(Fe,Co)2As2 films into iron pnictide 122-type. The
consequences for what the documents say:

  iron_chalcogenide_11   0.7687 Outcome A  ->  0.3737 Outcome B   (n 10 -> 12)
  iron_pnictide_122      0.3452 Outcome B  ->  0.4877 Outcome B   (n  9 -> 10)
  other_unclassified     0.9880 Outcome A  ->  the family dissolves

Two claims change rather than two numbers.

The diagnostic no longer returns three distinct regimes. It returns two: B for
both iron families and C for the MgB2 class. No family in the cohort now
requires sample-form conditioning outright, so "conditioning is essential" and
"conditioning is mandatory" become "conditioning is informative", and the
chalcogenide dispatch rule moves from Stage 2 to the same footing as the 122.

The abstract's 60% for the 122 family was stale twice over: it is the
pre-withdrawal 0.5988, not the 0.3452 the corrected cohort carried before this
fix, and it is 0.4877 now.

Every replacement is recomputed from data/phase_3_p31_variance_decomposition.csv
and the anchor table at run time. Refuses to write on any miss; treats an
already-applied edit as done.
"""
import argparse
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import docx  # noqa: E402
import pandas as pd  # noqa: E402

from apply_manuscript_edits import apply  # noqa: E402
from verify_deposit import aggregate_per_physical_sample  # noqa: E402

PLOTTED = ["iron_chalcogenide_11", "iron_pnictide_122", "conventional_AlB2"]


def numbers():
    vd = pd.read_csv(os.path.join(
        "data", "phase_3_p31_variance_decomposition.csv"))
    per = vd[vd.scope == "per_substructure"].set_index("substructure")
    a = pd.read_csv(os.path.join(
        "data", "phase_3_p31_jc_anchor_per_paper.csv"))
    agg = aggregate_per_physical_sample(a)
    ap = agg[agg.substructure.isin(PLOTTED)]
    r = {k: float(per.loc[k, "ratio_between_total"]) for k in PLOTTED}
    bands = sorted({("A" if v > 0.7 else ("B" if v >= 0.3 else "C"))
                    for v in r.values()})
    return dict(
        r=r, pct={k: round(100 * v) for k, v in r.items()},
        anchors=len(a), in_panels=int(a.substructure.isin(PLOTTED).sum()),
        markers=len(ap),
        per_fam={k: int((ap.substructure == k).sum()) for k in PLOTTED},
        n_bands=len(bands), bands=bands,
    )


def edits(N):
    ch, pn, mg = (N["r"]["iron_chalcogenide_11"], N["r"]["iron_pnictide_122"],
                  N["r"]["conventional_AlB2"])
    pc, pp, pm = (N["pct"]["iron_chalcogenide_11"],
                  N["pct"]["iron_pnictide_122"],
                  N["pct"]["conventional_AlB2"])
    word = {2: "two", 3: "three"}[N["n_bands"]]
    return dict(
        manuscript=[
            ("returns three distinct regimes across the studied families: "
             "sample form explains 73% of the within-family variance in the "
             "critical-current anchor for iron chalcogenide 11-type "
             "materials, 60% for iron pnictide 122-type, and 12% for "
             "MgB2-class.",
             "returns %s distinct regimes across the studied families: sample "
             "form explains %d%% of the within-family variance in the "
             "critical-current anchor for iron chalcogenide 11-type "
             "materials, %d%% for iron pnictide 122-type, and %d%% for "
             "MgB2-class." % (word, pc, pp, pm),
             "abstract; the 60% was the pre-withdrawal 122 value", None),
            ("returning three distinct regimes across the studied families",
             "returning %s distinct regimes across the studied families" % word,
             "conclusion", None),
            ("The diagnostic yields three distinct outcomes, summarized in "
             "Fig. 3. For iron chalcogenide 11-type materials the ratio is "
             "0.77: sample form accounts for most of the relevant variance, "
             "so sample-form conditioning is essential, the predictor uses "
             "Stage 2 medians, and the within-cell interquartile range "
             "captures the r",
             "The diagnostic yields %s distinct outcomes, summarized in "
             "Fig. 3. For iron chalcogenide 11-type materials the ratio is "
             "%.2f: sample form accounts for a substantial minority of the "
             "relevant variance, so sample-form conditioning is informative "
             "rather than required, the predictor uses Stage 2 medians where "
             "the cell is populated, and the within-cell interquartile range "
             "captures the r" % (word, ch),
             "Sec. III.A, the regime statement", None),
            ("Sample form explains 77%, 35%, and 12% of within-family anchor "
             "variance for iron chalcogenide 11-type, iron pnictide 122-type, "
             "and MgB2-class respectively.",
             "Sample form explains %d%%, %d%%, and %d%% of within-family "
             "anchor variance for iron chalcogenide 11-type, iron pnictide "
             "122-type, and MgB2-class respectively."
             % (pc, pp, pm), "figure callout", None),
            ("Of the 96 per-paper anchor records of Table I, 52 fall in the "
             "three families shown and the remainder in families not plotted "
             "here. The 52 collapse to the 34 markers drawn, because multiple "
             "isotherms of one physical sample are averaged into a single "
             "record before plotting: 10 for iron chalcogenide 11-type, 9 for "
             "iron pnictide 122-type, and 15 for MgB2-class.",
             "Of the %d per-paper anchor records of Table I, %d fall in the "
             "three families shown and the remainder in families not plotted "
             "here. The %d collapse to the %d markers drawn, because multiple "
             "isotherms of one physical sample are averaged into a single "
             "record before plotting: %d for iron chalcogenide 11-type, %d "
             "for iron pnictide 122-type, and %d for MgB2-class."
             % (N["anchors"], N["in_panels"], N["in_panels"], N["markers"],
                N["per_fam"]["iron_chalcogenide_11"],
                N["per_fam"]["iron_pnictide_122"],
                N["per_fam"]["conventional_AlB2"]),
             "Fig. 3 caption, recomputed", None),
            ("96 per-paper anchors in total, of which 52 fall in the three "
             "families plotted",
             "%d per-paper anchors in total, of which %d fall in the three "
             "families plotted" % (N["anchors"], N["in_panels"]),
             "same counts", None),
            ("For iron chalcogenide 11-type, sample-form conditioning is "
             "mandatory, so the dispatch would use the single-crystal "
             "empirical cell, the largest and best calibrated cell in the "
             "available literature.",
             "For iron chalcogenide 11-type, sample-form conditioning is "
             "informative but not decisive at a ratio of %.2f, so the dispatch "
             "would use the source-paper sample form where the cell is "
             "populated and fall back to the single-crystal cell, the largest "
             "and best calibrated cell in the available literature, "
             "otherwise." % ch,
             "no family now requires conditioning outright", None),
        ],
        letter=[
            ("The variance-decomposition test finds that sample form explains "
             "only 12% of the within-family variance in this class, against "
             "73% for the iron chalcogenides.",
             "The variance-decomposition test finds that sample form explains "
             "only %d%% of the within-family variance in this class, against "
             "%d%% for the iron chalcogenides." % (pm, pc),
             "recomputed after the classifier correction", None),
            ("The caption now gives both numbers, 52 of the 96 anchors "
             "falling in the three panels shown and the 34 markers they "
             "collapse to",
             "The caption now gives both numbers, %d of the %d anchors "
             "falling in the three panels shown and the %d markers they "
             "collapse to" % (N["in_panels"], N["anchors"], N["markers"]),
             "recomputed", None),
        ],
        supplement=[],
    )


_APOS = "['’]"


def _rewrap(find):
    return re.compile(r"\s+".join(
        "".join(_APOS if ch in "'’" else re.escape(ch) for ch in w)
        for w in find.split()))


def _flat(doc):
    parts = [p.text for p in doc.paragraphs]
    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return re.sub(r"\s+", " ", " ".join(parts))


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--ms", required=True)
    ap.add_argument("--supp", required=True)
    ap.add_argument("--letter", required=True)
    ap.add_argument("--letter-md", required=True)
    ap.add_argument("--out-dir", default=".")
    ap.add_argument("--dry-run", action="store_true")
    a = ap.parse_args()
    if not os.path.isdir("data"):
        sys.exit("run from the repository root")

    N = numbers()
    print("recomputed from the deposit\n")
    for k in PLOTTED:
        print("   %-22s ratio %.4f  %d%%  %d markers"
              % (k, N["r"][k], N["pct"][k], N["per_fam"][k]))
    print("   %d anchors, %d in the plotted families, %d markers, %d band(s): %s\n"
          % (N["anchors"], N["in_panels"], N["markers"], N["n_bands"],
             ", ".join(N["bands"])))

    E = edits(N)
    report, misses, applied = [], [], 0
    docs = {}
    for label, path in (("manuscript", a.ms), ("supplement", a.supp),
                        ("letter", a.letter)):
        if not E[label]:
            docs[label] = docx.Document(path)
            continue
        d = docx.Document(path)
        curly = "’" in _flat(d)
        E[label] = [((f.replace("'", "’") if curly else f), r, w, lo)
                    for f, r, w, lo in E[label]]
        sub = []
        apply(d, E[label], label, sub)
        text = _flat(d)
        for lb, done, find, repl, why in sub:
            if done:
                applied += 1
                report.append((lb, "ok", find, why))
            else:
                full = next(r for f, r, _, _ in E[label] if f[:72] == find)
                if _rewrap(full).search(text):
                    report.append((lb, "done", find, why))
                else:
                    misses.append(find)
                    report.append((lb, "MISS", find, why))
        docs[label] = d

    md = open(a.letter_md, encoding="utf-8").read()
    for find, repl, why, _ in E["letter"]:
        pat = _rewrap(find)
        if pat.search(md):
            md = pat.sub(lambda _m, r=repl: r, md, count=1)
            applied += 1
            report.append(("letter, markdown", "ok", find[:72], why))
        elif _rewrap(repl).search(re.sub(r"\s+", " ", md)):
            report.append(("letter, markdown", "done", find[:72], why))
        else:
            misses.append(find)
            report.append(("letter, markdown", "MISS", find[:72], why))

    for lb, state, find, why in report:
        print("   %-18s %-4s %s" % (lb, state, find))
        if state == "MISS":
            print("        %s" % why)
    print("\n   %d edit(s) applied this run\n" % applied)
    if misses:
        print("%d edit(s) missed; nothing written" % len(misses))
        return 1
    if a.dry_run:
        print("dry run, nothing written")
        return 0
    for label, path in (("manuscript", a.ms), ("supplement", a.supp),
                        ("letter", a.letter)):
        docs[label].save(os.path.join(a.out_dir, os.path.basename(path)))
    with open(os.path.join(a.out_dir, os.path.basename(a.letter_md)), "w",
              encoding="utf-8") as fh:
        fh.write(md)
    print("wrote all four artifacts to %s" % a.out_dir)
    return 0


if __name__ == "__main__":
    sys.exit(main())
